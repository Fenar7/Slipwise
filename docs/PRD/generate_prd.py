"""
SLIPWISE ONE — Full PRD Generator
Produces a 170+ page DOCX covering all phases, all sprints, all features.
Run: python generate_prd.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, datetime

OUTPUT = r"C:\Fenar\Zenxvio\Product-Works\Slipwise\Slipwise\docs\PRD\SLIPWISE_OS_EXPANSION_PRD_v2.docx"
doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1.0)
    section.top_margin  = section.bottom_margin = Inches(1.0)

# ── Style helpers ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x1B, 0x2A)
TEAL   = RGBColor(0x00, 0x87, 0x87)
DARK   = RGBColor(0x1C, 0x1C, 0x2E)
GRAY   = RGBColor(0x44, 0x44, 0x55)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF4, 0xF6, 0xF8)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),  "single")
        b.set(qn("w:sz"),   "4")
        b.set(qn("w:space"),"0")
        b.set(qn("w:color"),"008787")
        tcBorders.append(b)
    tcPr.append(tcBorders)

def h1(text, numbered=None):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(20)
    p.runs[0].font.bold = True
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(8)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = TEAL
    p.runs[0].font.size = Pt(15)
    p.runs[0].font.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = DARK
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.bold = True
    return p

def h4(text):
    p = doc.add_heading(text, level=4)
    p.runs[0].font.color.rgb = GRAY
    p.runs[0].font.size = Pt(12)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size  = Pt(10.5)
    p.runs[0].font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.runs[0].font.size = Pt(10.5)
    p.runs[0].font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(3)
    return p

def numbered(text, level=0):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.runs[0].font.size = Pt(10.5)
    p.runs[0].font.color.rgb = DARK
    return p

def divider():
    p = doc.add_paragraph("─" * 90)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xDD)
    p.runs[0].font.size = Pt(8)
    p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(4)

def page_break():
    doc.add_page_break()

def info_table(rows):
    """Renders a two-column key/value table."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,(k,v) in enumerate(rows):
        r = tbl.rows[i]
        r.cells[0].text = k
        r.cells[1].text = v
        r.cells[0].width = Inches(2.0)
        r.cells[1].width = Inches(4.5)
        set_cell_bg(r.cells[0], "0D1B2A")
        set_cell_bg(r.cells[1], "F4F6F8")
        for run in r.cells[0].paragraphs[0].runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(10)
        for run in r.cells[1].paragraphs[0].runs:
            run.font.color.rgb = DARK
            run.font.size = Pt(10)
        set_cell_border(r.cells[0])
        set_cell_border(r.cells[1])

def header_table(headers, rows_data, col_widths=None):
    """Renders a data table with a teal header row."""
    tbl = doc.add_table(rows=1+len(rows_data), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = tbl.rows[0]
    for j, h in enumerate(headers):
        cell = hrow.cells[j]
        cell.text = h
        if col_widths: cell.width = Inches(col_widths[j])
        set_cell_bg(cell, "008787")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(10)
    for i, row_data in enumerate(rows_data):
        row = tbl.rows[i+1]
        bg = "FFFFFF" if i % 2 == 0 else "F4F6F8"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(val)
            if col_widths: cell.width = Inches(col_widths[j])
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()

def sprint_header(phase, sprint_num, title, weeks, objective):
    doc.add_paragraph()
    divider()
    p = doc.add_paragraph()
    run = p.add_run(f"  PHASE {phase}  ·  SPRINT {sprint_num:02d}  ·  {title.upper()}  ·  {weeks}")
    run.font.bold  = True
    run.font.size  = Pt(11)
    run.font.color.rgb = WHITE
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),  "clear")
    shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"), "0D1B2A")
    p._p.get_or_add_pPr().append(shd)
    h3(f"Sprint Objective")
    body(objective)

def sprint_story(story_id, title, as_a, i_want, so_that, points, priority):
    h4(f"US-{story_id:03d} | {title}")
    info_table([
        ("As a…",     as_a),
        ("I want to…", i_want),
        ("So that…",  so_that),
        ("Story Points", str(points)),
        ("Priority",  priority),
    ])

def sprint_acceptance(criteria):
    h4("Acceptance Criteria")
    for c in criteria:
        bullet(c)

def sprint_tech(notes):
    h4("Technical Implementation Notes")
    for n in notes:
        bullet(n)


# ════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(72)
run = p.add_run("SLIPWISE ONE")
run.font.size  = Pt(42)
run.font.bold  = True
run.font.color.rgb = NAVY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p2 = doc.add_paragraph()
run2 = p2.add_run("Unified Business Operating System")
run2.font.size  = Pt(20)
run2.font.color.rgb = TEAL
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p3 = doc.add_paragraph()
run3 = p3.add_run("PRODUCT REQUIREMENTS DOCUMENT  ·  v2.0")
run3.font.size = Pt(14)
run3.font.bold = True
run3.font.color.rgb = GRAY
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

meta = [
    ("Document Status",  "FINAL — Engineering Handover"),
    ("Version",          "2.0"),
    ("Date",             datetime.date.today().strftime("%B %d, %Y")),
    ("Author",           "Slipwise — Office of the CTO"),
    ("Classification",   "INTERNAL — CONFIDENTIAL"),
    ("Phases Covered",   "Phase A · Phase B · Phase C"),
    ("Total Sprints",    "38 Sprints"),
    ("Est. Duration",    "19 Months"),
    ("Target Stack",     "Next.js 14 · Prisma ORM · PostgreSQL · Supabase Auth · AWS"),
    ("Audience",         "Engineering Leads, Product Managers, QA, DevOps"),
]
info_table(meta)
page_break()


# ════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════════════
h1("Table of Contents")

toc_items = [
    ("1",  "Executive Summary",                                    "5"),
    ("2",  "Product Vision and Strategic Context",                 "7"),
    ("3",  "Glossary of Terms",                                    "10"),
    ("4",  "Technology Architecture",                              "14"),
    ("5",  "Git Strategy and Branch Model",                        "18"),
    ("6",  "Security and Compliance Mandates",                     "21"),
    ("7",  "AWS Infrastructure Design",                            "25"),
    ("8",  "Phase A — Foundation (Sprints 1–9)",                   "30"),
    ("8.1","Sprint 01 — Org Tree and Entity Hierarchy",            "31"),
    ("8.2","Sprint 02 — RBAC Engine and Permission Matrix",        "34"),
    ("8.3","Sprint 03 — Member Onboarding Wizard",                 "38"),
    ("8.4","Sprint 04 — Work OS Data Models",                      "42"),
    ("8.5","Sprint 05 — Task CRUD and List View",                  "46"),
    ("8.6","Sprint 06 — Board View, Dependencies, Real-Time",      "50"),
    ("8.7","Sprint 07 — Employee Profile and Directory",           "54"),
    ("8.8","Sprint 08 — Leave Management Engine",                  "58"),
    ("8.9","Sprint 09 — Attendance, Holidays, Team Calendar",      "62"),
    ("9",  "Phase B — Feature Depth (Sprints 10–23)",              "67"),
    ("9.1","Sprint 10 — Time Tracking",                            "68"),
    ("9.2","Sprint 11 — Sprint Management and Backlog",            "72"),
    ("9.3","Sprint 12 — Gantt Chart and Mind Map Views",           "76"),
    ("9.4","Sprint 13 — Workload, Forms, Templates",               "80"),
    ("9.5","Sprint 14 — Performance Review Cycles",                "84"),
    ("9.6","Sprint 15 — Onboarding and Offboarding Automation",    "88"),
    ("9.7","Sprint 16 — Payroll Integration Layer",                "92"),
    ("9.8","Sprint 17 — CRM Contacts and Companies",               "96"),
    ("9.9","Sprint 18 — Pipelines, Deals, and Revenue Tracking",   "100"),
    ("9.10","Sprint 19 — Lead Scoring and Email Sequences",        "104"),
    ("9.11","Sprint 20 — Meeting Scheduling and Activity Log",     "108"),
    ("9.12","Sprint 21 — ITSM Service Catalogue and Ticket Submit","112"),
    ("9.13","Sprint 22 — Ticket Routing, SLA, and Escalation",    "116"),
    ("9.14","Sprint 23 — CSAT, Reporting, and Knowledge Base",     "120"),
    ("10", "Phase C — Intelligence and Scale (Sprints 24–38)",     "125"),
    ("10.1","Sprint 24 — CX Health Scoring and Onboarding Plays",  "126"),
    ("10.2","Sprint 25 — Goals and OKR Engine",                    "130"),
    ("10.3","Sprint 26 — Dashboards and Analytics Studio",         "134"),
    ("10.4","Sprint 27 — Docs, Wiki, and Knowledge Management",    "138"),
    ("10.5","Sprint 28 — Advanced Docs: Embeds and Collaboration", "141"),
    ("10.6","Sprint 29 — Visual Automation Builder",               "144"),
    ("10.7","Sprint 30 — Trigger Library and Integration Hooks",   "148"),
    ("10.8","Sprint 31 — Zapier/Webhook External Automation",      "152"),
    ("10.9","Sprint 32 — AI Writing and Summarisation Layer",      "155"),
    ("10.10","Sprint 33 — AI Predictive Analytics and Insights",   "158"),
    ("10.11","Sprint 34 — Google Workspace Integration",           "161"),
    ("10.12","Sprint 35 — Slack, Zoom, and Calendar Sync",         "164"),
    ("10.13","Sprint 36 — Mobile PWA Foundation",                  "167"),
    ("10.14","Sprint 37 — Mobile Offline and Push Notifications",  "170"),
    ("10.15","Sprint 38 — Hardening, Audit, and Public Launch",    "173"),
    ("11", "Database Schema Master Reference",                     "177"),
    ("12", "API Endpoint Index",                                   "185"),
    ("13", "Testing Strategy",                                     "191"),
    ("14", "Master Development Timeline",                          "196"),
    ("15", "Risk Register",                                        "202"),
    ("16", "Appendix",                                             "206"),
]

tbl = doc.add_table(rows=len(toc_items), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i,(num,title,pg) in enumerate(toc_items):
    row = tbl.rows[i]
    row.cells[0].text = num
    row.cells[1].text = title
    row.cells[2].text = pg
    row.cells[0].width = Inches(0.6)
    row.cells[1].width = Inches(5.5)
    row.cells[2].width = Inches(0.8)
    bg = "F4F6F8" if i % 2 == 0 else "FFFFFF"
    for c in row.cells:
        set_cell_bg(c, bg)
        for run in c.paragraphs[0].runs:
            run.font.size = Pt(10)
page_break()


# ════════════════════════════════════════════════════════════════════════════
#  §1 — EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════════
h1("1. Executive Summary")
body("Slipwise One is a single, unified Business Operating System built for modern companies who are tired of managing six to twelve different SaaS tools that never talk to each other. Instead of paying separate subscriptions for a project management tool, an HR platform, a CRM, a customer success platform, an IT service desk, and a document wiki, Slipwise brings every operational surface into one coherent product with a single permission model, a single notification layer, and a single source of truth for every entity — Person, Task, Deal, Ticket, Document — in the business.")
body("This document is the formal Product Requirements Document for the Slipwise One expansion. It covers nineteen months of engineering work spread across thirty-eight two-week sprints, organised into three phases. Every sprint in this document is described at the story level: each user story has a title, a full As-a / I-want / So-that breakdown, story point estimates, acceptance criteria, and technical implementation notes written by the CTO's office so that engineering teams can begin work immediately without requiring further specification sessions for the features covered.")
body("The first phase, called Phase A, spans Sprints 01 through 09 and lays the architectural foundation. This phase builds the Organisation Tree and Entity Hierarchy that defines how a company is structured inside Slipwise, the Role-Based Access Control engine and granular Permission Matrix that controls what each employee can see and do, the Member Onboarding Wizard that walks administrators through inviting and configuring new team members, the Work OS data models and views that replace tools like ClickUp and Asana, and the core Human Resources modules covering employee profiles, leave management, attendance tracking, and team calendars.")
body("The second phase, called Phase B, spans Sprints 10 through 23 and adds depth and breadth to the platform. It introduces Time Tracking, Sprint and Backlog management for engineering teams, Gantt Charts and Mind Maps as additional Work OS views, Workload and Capacity planning, Performance Review cycles, automated Onboarding and Offboarding workflows, a Payroll integration layer, a full CRM with contact and company management, deal pipelines, lead scoring, email sequences, and meeting scheduling. Phase B also introduces the ITSM (Internal Ticketing) module with a full service catalogue, ticket routing with SLA enforcement, escalation rules, and CSAT collection.")
body("The third phase, called Phase C, spans Sprints 24 through 38 and introduces intelligence and scale. It builds the Customer Experience module with health scoring and onboarding playbooks, adds an OKR and Goals engine, a configurable Analytics Studio, a full Docs and Wiki system with collaborative real-time editing, a no-code Visual Automation Builder, an AI writing and summarisation layer, predictive analytics, integrations with Google Workspace, Slack, Zoom, and Calendar platforms, a Mobile PWA, and a comprehensive hardening, security audit, and public launch sprint.")

h2("1.1 Business Case")
body("Research published by Gartner and McKinsey consistently shows that the average knowledge-work employee switches between nine different applications in a single working day. This context-switching costs an estimated 23 minutes of deep focus per switch. For a 50-person company, this translates to thousands of hours of lost productivity per year. Beyond productivity loss, data fragmentation across tools creates reporting blind spots, delays in decision-making, and substantial SaaS spend that scales poorly.")
body("Slipwise's consolidated model solves this. By unifying data under one permission model and one data graph, every manager in the company can answer a question like 'Show me everything related to this customer — the deal, the onboarding tasks, the support tickets, and the health score' without opening four different tabs. This is the core value proposition and every engineering decision in this document is in service of that goal.")

h2("1.2 Success Metrics (KPIs for V1 Launch)")
header_table(
    ["KPI", "Baseline", "Target at 6 Months", "Target at 12 Months"],
    [
        ["Daily Active Users (DAU / WAU)", "—", "≥ 60%", "≥ 70%"],
        ["Avg. modules used per session", "—", "≥ 3.5", "≥ 5"],
        ["Task creation per user per week", "—", "≥ 8", "≥ 12"],
        ["Net Promoter Score (NPS)", "—", "≥ 40", "≥ 55"],
        ["P95 API response time", "—", "< 250 ms", "< 150 ms"],
        ["System uptime (monthly)", "—", "99.5%", "99.9%"],
        ["Leave approval cycle time", "—", "< 4 hours", "< 2 hours"],
        ["CRM deal stage conversion rate", "—", "Baseline set", "+10% MoM"],
        ["ITSM ticket first-response SLA", "—", "≥ 90%", "≥ 95%"],
        ["Customer health score accuracy", "—", "—", "Within ±10%"],
    ],
    col_widths=[2.2, 1.3, 1.8, 1.8]
)
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §2 — PRODUCT VISION
# ════════════════════════════════════════════════════════════════════════════
h1("2. Product Vision and Strategic Context")
h2("2.1 Vision Statement")
body("Slipwise One will be the operating system of record for every company that uses it. Just as a smartphone operating system provides the foundation that every application runs on top of, Slipwise provides the data foundation, the permission model, the identity layer, and the communication fabric that every internal workflow runs on top of. When a company runs on Slipwise, there is no need for a separate task manager, HR tool, CRM, wiki, service desk, or analytics platform. Everything is connected, everything is permissioned correctly, and every employee sees only what they are authorised to see.")

h2("2.2 The Ten Product Pillars")
pillars = [
    ("1. People Management (HRIS)", "A complete Human Resources Information System covering organisation hierarchy, role definitions, employee profiles, leave management, attendance, payroll integration, and performance reviews. Built to replace BambooHR or Rippling for companies under 1,000 employees."),
    ("2. Work OS", "A hierarchical project and task management system covering Workspaces, Spaces, Folders, Lists, Tasks, and Subtasks with ten different views (List, Board, Gantt, Calendar, Workload, Timeline, Mind Map, Table, Dashboard, Form). Built to replace ClickUp or Asana."),
    ("3. Role-Based Access Control", "A granular, per-resource permission system that allows administrators to control access at the module, sub-module, mailbox, and document level for every user in the organisation, with support for role templates and inherited permissions."),
    ("4. CRM and BDR Engine", "A full B2B sales pipeline covering Contacts, Companies, Deals, Activities, Lead Scoring, Email Sequences, and Revenue Forecasting. Built to compete with HubSpot Sales Hub for SMB companies."),
    ("5. Customer Experience (CX)", "A post-sales customer success platform covering Customer Health Scoring, Onboarding Playbooks, Renewal Pipelines, and CSAT measurement, designed to reduce churn and increase Net Revenue Retention (NRR)."),
    ("6. Internal Service Desk (ITSM)", "A Jira Service Management or Zendesk-equivalent for internal teams: service catalogues, ticket routing, SLA enforcement, escalation trees, CSAT, and a knowledge base with resolution suggestion."),
    ("7. Docs and Wiki", "A collaborative, block-based document editor that supports rich media embeds, version history, nested pages, templates, public sharing, and real-time multi-user cursors."),
    ("8. Visual Automation Builder", "A no-code trigger-action workflow engine that allows power users to automate multi-step processes across any module — for example, 'When a deal is won, create an onboarding task list and notify the CS team.'"),
    ("9. AI Layer", "An embedded AI assistant for writing, summarisation, task generation, lead scoring, health score prediction, and sentiment analysis — all powered by a secure LLM integration that never exposes company data to model training."),
    ("10. Analytics Studio", "A configurable business intelligence layer that generates dashboards, reports, and charts from any data in the platform — across HR, CRM, CX, ITSM, and Work OS — with scheduled report delivery and export."),
]
for name, desc in pillars:
    h3(name)
    body(desc)

h2("2.3 Competitive Landscape")
header_table(
    ["Platform", "Category", "Slipwise Advantage"],
    [
        ["ClickUp",          "Work OS",        "Slipwise adds HRIS, CRM, CX, ITSM natively — ClickUp has none of these."],
        ["Asana",            "Work OS",        "Slipwise has more views and native automation without Zapier dependency."],
        ["HubSpot Sales Hub","CRM",            "Slipwise is cheaper for SMBs and natively connected to task and project data."],
        ["BambooHR",         "HRIS",           "Slipwise's HRIS is embedded in the operating system, not a silo."],
        ["Jira Service Mgmt","ITSM",           "Slipwise ITSM is simpler to configure and connected to the same user identity."],
        ["Notion",           "Docs / Wiki",    "Slipwise Docs is embedded in context (tasks, deals) not a separate tool."],
        ["Gainsight",        "Customer Success","Slipwise CX is included at no extra cost vs. Gainsight's $30k+ ACV."],
        ["Tableau",          "Analytics",      "Slipwise Analytics is built on live operational data with zero ETL needed."],
    ],
    col_widths=[1.6, 1.4, 4.1]
)
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §3 — GLOSSARY
# ════════════════════════════════════════════════════════════════════════════
h1("3. Glossary of Terms")
body("This section defines every term used in this document. Engineering team members should read this section before reading the sprint specifications to ensure shared vocabulary across functions.")

glossary = [
    ("RBAC", "Role-Based Access Control. A method of restricting system access based on the roles assigned to individual users within an organisation. In Slipwise, RBAC is implemented at a very granular level: permissions can be granted or denied for individual modules, sub-features, or even specific mailboxes or document folders."),
    ("Granular Permission Grant", "A database record in the GranularPermissionGrant table that explicitly allows or denies a specific permission (e.g., 'can_view_invoices') to a specific user within a specific organisation. The absence of a grant is treated as denial by default."),
    ("Workspace", "The top-level container in Slipwise. A single company (tenant) has one Workspace. All Spaces, Members, Deals, Documents, and Tickets belong to a Workspace."),
    ("Space", "The first tier of the Work OS hierarchy below Workspace. Spaces typically represent departments or major project areas (e.g., 'Engineering', 'Marketing', 'Client Projects')."),
    ("Folder", "An optional second tier within a Space used to group related Lists (e.g., 'Q3 Campaigns' inside the 'Marketing' Space)."),
    ("List", "A collection of Tasks. Corresponds to a project or a recurring work queue. Lists can exist inside Folders or directly inside Spaces."),
    ("Task", "The atomic unit of work in the Work OS. Tasks have an assignee, due date, priority, status, time estimate, and can have Subtasks, Comments, Attachments, and Custom Fields."),
    ("Subtask", "A child task nested under a parent Task. Subtasks have their own assignees, due dates, and statuses independent of the parent."),
    ("Custom Status", "A named status value defined by the organisation (e.g., 'In Review', 'Blocked', 'Pending Client'). Custom statuses belong to a Space and are colour-coded."),
    ("Sprint", "A fixed time-box, typically two weeks, in which an engineering team commits to completing a set of user stories. Sprints are distinct from agile Sprints in Work OS."),
    ("Backlog", "The full, unprioritised list of tasks and user stories that exist for a project but have not yet been assigned to a Sprint."),
    ("Gantt Chart", "A horizontal bar chart view that shows tasks on a timeline with start and end dates, dependencies between tasks shown as connecting lines, and critical path highlighting."),
    ("Workload View", "A calendar-style view that shows how much work each team member is assigned across a time range, colour-coded by capacity utilisation."),
    ("OKR", "Objectives and Key Results. A goal-setting framework where an Objective is a qualitative, aspirational direction and Key Results are quantitative, measurable outcomes that define success for that Objective."),
    ("NRR", "Net Revenue Retention. A CX metric calculated as (Starting MRR + Expansion MRR - Churn MRR - Contraction MRR) / Starting MRR × 100. An NRR above 100% means the company is growing revenue from existing customers."),
    ("CSAT", "Customer Satisfaction Score. A survey metric, usually a 1–5 or 1–10 rating, collected after a support interaction or service delivery event to measure satisfaction."),
    ("NPS", "Net Promoter Score. A loyalty metric asking customers 'How likely are you to recommend us?' on a 0–10 scale. Promoters (9–10) minus Detractors (0–6) equals the NPS."),
    ("SLA", "Service Level Agreement. A commitment about response or resolution time. In ITSM, an SLA might state that P1 tickets must receive a first response within 1 hour."),
    ("Escalation", "The automatic or manual process of routing an unresolved ticket to a higher-priority or different support tier when SLA thresholds are breached."),
    ("Pipeline", "In CRM, a visual, stage-based workflow representing the journey of a deal from initial contact to close. Each stage has entry criteria and a probability percentage."),
    ("Lead Scoring", "An algorithmic system that assigns a numerical score to a lead based on firmographic data, behavioural signals (email opens, page visits), and CRM activity, indicating how likely the lead is to convert."),
    ("Health Score", "In CX, a composite score for a customer account calculated from usage metrics, support ticket volume, NPS responses, and renewal proximity. Low health scores trigger automated alerts."),
    ("Playbook", "In CX, a structured, milestone-based workflow assigned to a customer account to guide them through onboarding, expansion, or renewal. Each milestone has tasks, templates, and deadlines."),
    ("BDR", "Business Development Representative. A sales role focused on outbound prospecting and lead qualification. The BDR module covers the tools that BDR teams use daily."),
    ("PWA", "Progressive Web App. A web application built with modern browser APIs to behave like a native mobile application — supporting offline access, push notifications, and home-screen installation — without requiring submission to an app store."),
    ("HRIS", "Human Resource Information System. The database and set of processes that manage employee records, compensation, benefits, leave balances, and organisational structure."),
    ("ETL", "Extract, Transform, Load. A data pipeline process. Slipwise Analytics Studio avoids ETL by reading directly from the operational PostgreSQL database via optimised read replicas."),
    ("Supabase Auth", "The authentication layer used by Slipwise, built on top of GoTrue. It provides JWT-based sessions, magic link login, OAuth providers, and row-level security integration with PostgreSQL."),
    ("Prisma ORM", "The Object-Relational Mapper used by Slipwise to interact with PostgreSQL. Prisma provides type-safe database queries and manages schema migrations."),
    ("Middleware Guard", "A Next.js middleware function that runs at the edge before any page or API route is served. It validates the user's JWT, resolves their organisation, and checks their RBAC permissions before allowing access."),
]

for term, definition in glossary:
    h3(term)
    body(definition)
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §4 — TECHNOLOGY ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════════
h1("4. Technology Architecture")
h2("4.1 Architecture Philosophy")
body("Slipwise is built as a server-rendered, full-stack Next.js application deployed on AWS infrastructure. The choice of Next.js 14 with the App Router is deliberate: server components reduce the amount of JavaScript sent to the browser, server actions allow form submissions and mutations without a separate REST API layer for simple cases, and the file-system-based routing makes the codebase navigable for engineers who join the team at any phase. Complex queries and batch operations are exposed as REST API routes under /api/v1/, which can be consumed by both the Next.js frontend and future mobile clients.")
body("The database is PostgreSQL 15, accessed through Prisma ORM for type safety and schema management. All multi-tenant data is partitioned by organisationId with indexed foreign keys. Authentication is handled by Supabase Auth, providing JWT sessions with a 24-hour expiry and a 7-day refresh window. Row-Level Security (RLS) policies in PostgreSQL provide a second layer of protection at the database level, ensuring that even if an API query accidentally omits an organisationId filter, the database itself will return only the requesting organisation's data.")
body("Real-time collaboration features — such as live task updates, document cursor positions, and notification delivery — are powered by a combination of Supabase Realtime (for simple presence and record change events) and a Redis Pub/Sub layer (for high-volume events in the task and comment systems). The Redis layer is hosted on AWS ElastiCache in a cluster mode configuration for horizontal scalability.")

h2("4.2 Technology Stack Reference")
header_table(
    ["Layer", "Technology", "Version", "Purpose"],
    [
        ["Frontend Framework", "Next.js",         "14.x",  "Full-stack React framework, App Router, Server Components"],
        ["Language",          "TypeScript",        "5.x",   "Type safety across frontend and backend"],
        ["Database",          "PostgreSQL",        "15.x",  "Primary relational database, all operational data"],
        ["ORM",               "Prisma",            "5.x",   "Type-safe DB access, migration management, schema introspection"],
        ["Authentication",    "Supabase Auth",     "2.x",   "JWT sessions, OAuth, magic links, MFA"],
        ["Real-Time",         "Supabase Realtime", "2.x",   "Presence, record change subscriptions (low-volume)"],
        ["Real-Time (high)",  "Redis Pub/Sub",     "7.x",   "High-volume event distribution (tasks, comments, notifications)"],
        ["Cache",             "Redis",             "7.x",   "Session caching, rate limiting, background job queues"],
        ["File Storage",      "AWS S3",            "—",     "Document attachments, avatar images, export files"],
        ["CDN",               "AWS CloudFront",    "—",     "Edge caching of S3 assets and static Next.js output"],
        ["Background Jobs",   "BullMQ",            "4.x",   "Email dispatch, automation triggers, report generation"],
        ["Email Sending",     "AWS SES",           "v2",    "Transactional and sequence emails"],
        ["Search",            "PostgreSQL FTS",    "—",     "Full-text search on tasks, docs, contacts via tsvector"],
        ["AI / LLM",          "Google Gemini API", "1.5",   "Writing, summarisation, scoring, predictions"],
        ["Monitoring",        "AWS CloudWatch",    "—",     "Logs, metrics, alarms, distributed tracing"],
        ["CI/CD",             "GitHub Actions",    "—",     "Lint, test, build, deploy pipeline"],
        ["Containerisation",  "Docker",            "24.x",  "Development environment and production image"],
        ["Orchestration",     "AWS ECS Fargate",   "—",     "Serverless container execution"],
        ["DNS / WAF",         "AWS Route 53 + WAF","—",     "Domain management, DDoS protection, IP filtering"],
        ["Secrets",           "AWS Secrets Manager","—",    "API keys, DB credentials, never in .env files"],
    ],
    col_widths=[1.6, 1.6, 0.7, 3.2]
)

h2("4.3 Application Directory Structure")
body("The Slipwise codebase follows a feature-co-located structure within the Next.js App Router. Every module (work-os, hris, crm, itsm, cx, docs, automation, ai) has its own directory under src/app/(app)/[module]/ containing its pages, components, and module-specific API routes. Shared utilities, hooks, and types live in src/lib/, src/hooks/, and src/types/ respectively. The Prisma schema is a single file at prisma/schema.prisma, but the migration history is organised with descriptive names that map to sprint numbers.")
body("This structure ensures that a developer working on the CRM module does not need to understand the HRIS module's internal structure — their code is co-located and isolated. However, because they share the same Prisma client and permission middleware, cross-module data access is always type-safe and always permission-checked.")
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §5 — GIT STRATEGY
# ════════════════════════════════════════════════════════════════════════════
h1("5. Git Strategy and Branch Model")
h2("5.1 Branch Topology")
body("Slipwise uses a trunk-based development model with short-lived feature branches and a protected main branch. The main branch always represents the current production-deployable state of the codebase. Feature branches are created from main, kept alive for a maximum of five working days (one sprint partial), and merged back via Pull Request after passing all automated checks.")
body("Because many Phase B and Phase C features can be built in parallel — for example, CRM development can run concurrently with HRIS Performance Reviews — the repository uses Git Worktrees to allow parallel development in the same repository without duplicating the full codebase on disk. Each worktree is bound to a long-lived integration branch for its phase, and short-lived feature branches are created inside that worktree.")

h2("5.2 Branch Naming Convention")
header_table(
    ["Branch Type",     "Pattern",                        "Example"],
    [
        ["Main (production)",   "main",                       "main"],
        ["Phase integration",   "phase/{phase-letter}",       "phase/b"],
        ["Feature",             "feat/{sprint-num}/{slug}",   "feat/10/time-tracking"],
        ["Bug fix",             "fix/{ticket-id}/{slug}",     "fix/SLW-412/leave-overlap"],
        ["Hotfix (prod)",       "hotfix/{slug}",              "hotfix/permission-bypass"],
        ["Release candidate",   "release/{version}",         "release/1.3.0"],
        ["Documentation",       "docs/{slug}",                "docs/api-reference-update"],
        ["Chore / DevOps",      "chore/{slug}",              "chore/upgrade-prisma-5"],
    ],
    col_widths=[1.8, 2.8, 2.5]
)

h2("5.3 Pull Request Requirements")
body("No code is merged to main or a phase integration branch without satisfying every item in the PR checklist. This checklist is enforced by a GitHub Actions PR template and a required status check that validates the checklist is complete.")
checklist = [
    "All automated unit tests pass (0 failures, 0 skipped).",
    "All integration tests for the affected module pass.",
    "TypeScript compilation succeeds with zero errors (strict mode).",
    "ESLint reports zero errors (warnings are allowed, max 5 per PR).",
    "Prisma migrations are included and have been tested against a fresh seed database.",
    "API endpoint changes are reflected in the OpenAPI spec file.",
    "At least one peer reviewer (who did not write the code) has approved.",
    "No PR description fields are left blank (summary, testing steps, screenshots/recordings for UI changes).",
    "Bundle size increase is justified in the PR description if it exceeds 10KB gzipped.",
    "All new environment variables are documented in .env.example and in AWS Secrets Manager.",
]
for item in checklist:
    bullet(item)

h2("5.4 Commit Message Convention")
body("Slipwise follows the Conventional Commits specification (v1.0.0). Every commit message must begin with a type prefix followed by an optional scope in parentheses, a colon, a space, and then the imperative-mood description. The types used are: feat (new feature), fix (bug fix), docs (documentation), chore (tooling or configuration), refactor (code restructuring without behaviour change), test (adding or modifying tests), and perf (performance improvement). Examples: 'feat(crm): add lead scoring algorithm', 'fix(hris): correct leave balance calculation on leap year', 'perf(work-os): add tsvector index to task description column'.")
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §6 — SECURITY MANDATES
# ════════════════════════════════════════════════════════════════════════════
h1("6. Security and Compliance Mandates")
body("Every engineer on the Slipwise team is responsible for security. Security is not a phase at the end of development — it is a gate at every PR review. The following mandates are non-negotiable. Violations are grounds for PR rejection and must be remediated before merge.")

security_mandates = [
    ("M-001: Multi-Tenancy Isolation", "Every single database query that touches organisation-scoped data MUST include an explicit organisationId filter. This is not optional and will not be assumed to be safe without the filter. To prevent accidental omission, all service-layer functions that retrieve organisation-scoped data must accept an organisationId parameter as their first argument and throw a TypeScript compile-time error if it is omitted. Code reviewers must verify this for every new database query introduced in a PR."),
    ("M-002: Permission Check Before Data Access", "Every API route handler must call the checkPermission(userId, organisationId, permissionKey) function before returning any data or performing any mutation. This function resolves the user's GranularPermissionGrant records and throws a 403 Forbidden error if the permission is not granted. Bypassing this check for convenience or speed is strictly prohibited."),
    ("M-003: Input Validation with Zod", "All incoming data from HTTP request bodies, query strings, and route parameters must be validated using a Zod schema before being processed. Prisma's type safety does not substitute for input validation because Prisma does not validate business constraints such as maximum string length, allowed enum values, or non-negative numbers. Every API route must define and use a corresponding Zod schema."),
    ("M-004: No Secrets in Code or Environment Files", "API keys, database connection strings, JWT secrets, and third-party service credentials must never appear in source code, .env files committed to version control, or PR descriptions. All secrets are stored in AWS Secrets Manager and injected as environment variables at container startup by ECS task definitions. The .env.example file contains only placeholder values and is the only secrets-related file committed to the repository."),
    ("M-005: JWT Validation on Every API Route", "The Supabase Auth JWT must be validated on every API route using the verifyJWT middleware. This middleware checks the token's signature, expiry, and issuer. Trusting a userId from a request body or query string without a corresponding valid JWT is a critical security vulnerability and must never occur in Slipwise code."),
    ("M-006: Rate Limiting", "All public-facing API routes (those that can be called without authentication) must be rate-limited using the Redis-backed rate limiter middleware. Authenticated routes must also be rate-limited at a higher threshold to prevent abuse. Rate limit thresholds are defined in the rateLimit.config.ts file and must be reviewed for any new route."),
    ("M-007: SQL Injection Prevention", "Because Slipwise uses Prisma ORM for all database access, raw SQL is forbidden except in explicitly reviewed migration scripts or read-model queries that cannot be expressed with Prisma. Any use of Prisma's $queryRaw or $executeRaw must be reviewed by a senior engineer and the query must use parameterised placeholders — never string interpolation."),
    ("M-008: CSRF Protection", "All state-mutating API routes (POST, PUT, PATCH, DELETE) must include CSRF token validation. Next.js server actions provide built-in CSRF protection, but custom API route handlers under /api/ must use the csrfProtect middleware."),
    ("M-009: Content Security Policy", "The Next.js application must set a strict Content Security Policy header on all responses. The CSP must prohibit inline scripts (no unsafe-inline), restrict script sources to the application's own domain and explicitly whitelisted CDN domains, and block all frame-ancestors except the application itself. The CSP configuration lives in next.config.js and must be reviewed when new third-party scripts are added."),
    ("M-010: File Upload Validation", "All file upload endpoints must validate the file's MIME type, maximum size (10 MB default, 50 MB for document attachments), and file extension before uploading to S3. Files must be stored in a private S3 bucket and accessed via pre-signed URLs with a 1-hour expiry. Direct S3 URLs must never be exposed to clients."),
    ("M-011: Audit Logging", "Every state-mutating operation on sensitive entities (users, permissions, deals, documents, payroll records) must create a record in the AuditLog table. The audit log record includes the actor's userId, the action performed, the entity type and ID affected, the timestamp, and the actor's IP address. Audit logs are immutable — no API route may delete audit log records."),
    ("M-012: Data Encryption at Rest", "All PostgreSQL data at rest is encrypted using AWS RDS's AES-256 encryption. S3 buckets use SSE-S3 encryption by default. Particularly sensitive fields — such as payroll records and national ID numbers — are additionally encrypted at the application layer using a field-level encryption utility before being stored, so that even a database administrator with direct DB access cannot read raw sensitive values."),
    ("M-013: Dependency Vulnerability Scanning", "The GitHub Actions CI pipeline must include a step that runs npm audit and Snyk scan on every PR. PRs that introduce new dependencies with High or Critical CVEs must not be merged until the vulnerability is resolved or a written exception is approved by the engineering lead."),
    ("M-014: Security Review Gate for New Modules", "Before any new module (CRM, ITSM, AI, etc.) is merged to the phase integration branch, it must undergo a focused security review session. The review must cover the module's permission model, data access patterns, API surface, and any third-party integrations it introduces. The review outcomes must be documented and linked from the PR."),
]

for code, description in security_mandates:
    h3(f"{code}")
    body(description)
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §7 — AWS INFRASTRUCTURE
# ════════════════════════════════════════════════════════════════════════════
h1("7. AWS Infrastructure Design")
h2("7.1 Infrastructure Overview")
body("Slipwise's production infrastructure runs entirely on Amazon Web Services (AWS) in the ap-south-1 (Mumbai) primary region with a warm standby in ap-southeast-1 (Singapore) for disaster recovery. All infrastructure is defined as code using AWS CloudFormation templates, versioned in a separate private repository (slipwise-infra), and applied through a GitHub Actions workflow that requires manual approval for production changes.")

h2("7.2 Service Architecture")
services = [
    ("AWS ECS Fargate",         "Runs the Next.js application container. Auto-scales between 2 and 20 tasks based on CPU and memory metrics. Each task has 1 vCPU and 2 GB RAM for the web tier. Background worker tasks (BullMQ consumers) run in a separate Fargate service with their own auto-scaling policy."),
    ("AWS RDS PostgreSQL",      "Multi-AZ PostgreSQL 15 cluster. Primary instance is db.r6g.large (2 vCPU, 16 GB RAM). A read replica in the same AZ handles Analytics Studio queries to prevent read pressure on the primary. Point-in-time recovery is enabled with a 35-day retention window. Automated daily snapshots are retained for 7 days."),
    ("AWS ElastiCache Redis",   "Redis 7 in cluster mode with 3 shards and 1 replica per shard, providing 18 GB total cache memory. Used for BullMQ job queues, rate limiting counters, session data, and real-time Pub/Sub channels."),
    ("AWS S3",                  "Three buckets: (1) slipwise-uploads-prod for all user-uploaded files, (2) slipwise-exports-prod for generated reports and document exports, (3) slipwise-backups-prod for database backup archives. All buckets are private with public access blocked. Versioning enabled on uploads and backups buckets."),
    ("AWS CloudFront",          "CDN distribution in front of S3 for static assets and signed URL generation for file downloads. Separate distribution for Next.js ISR pages with a 60-second TTL on cached pages. WAF is attached to the CloudFront distribution with AWS Managed Rules enabled plus custom rules for known bot signatures."),
    ("AWS SES",                 "Simple Email Service for all outgoing email. Two sending identities: noreply@slipwise.com for transactional emails and sequences@slipwise.com for CRM email sequences. Bounce and complaint handling via SNS topics that automatically suppress problematic addresses."),
    ("AWS Secrets Manager",     "Stores all credentials and API keys. ECS task definitions reference Secrets Manager ARNs for environment injection. Secret rotation is configured for database passwords (90-day rotation) and JWT secrets (180-day rotation)."),
    ("AWS CloudWatch",          "Log groups for each ECS service with 30-day retention. Custom metrics for business KPIs (tasks created, tickets resolved, deals won) published from the application. Alarms on P99 API latency > 1s, error rate > 1%, and ECS CPU > 80% sustained for 5 minutes."),
    ("AWS Route 53",            "DNS for *.slipwise.com with health check-based routing that fails over to the Singapore standby region if the Mumbai region health check fails for 3 consecutive 10-second intervals."),
]
for service, description in services:
    h3(service)
    body(description)

h2("7.3 Network Security")
body("The production VPC uses a three-tier subnet model. Public subnets contain only the Application Load Balancer and NAT Gateway. The application tier (ECS Fargate tasks) runs in private subnets with outbound-only internet access through the NAT Gateway. The data tier (RDS and ElastiCache) runs in isolated subnets with no internet access whatsoever — they accept connections only from the application tier security group.")
body("Security Group rules follow the principle of least privilege. The ALB security group accepts HTTPS (443) from 0.0.0.0/0. The application security group accepts traffic only from the ALB security group on port 3000. The RDS security group accepts traffic only from the application security group on port 5432. ElastiCache accepts traffic only from the application security group on port 6379.")
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §8 — PHASE A (SPRINTS 01–09)
# ════════════════════════════════════════════════════════════════════════════
h1("8. Phase A — Foundation (Sprints 01–09)")
body("Phase A establishes the non-negotiable foundation that every subsequent module depends on. No Phase B or Phase C work can begin until Phase A is complete, tested, and deployed to the staging environment. Phase A covers the organisational data model, the permission system, the member onboarding experience, the core Work OS data layer and its first two views, and the foundational HRIS modules. Phase A spans Sprints 01 through 09, representing nine two-week sprints or approximately 18 weeks of development.")

header_table(
    ["Sprint", "Title", "Weeks", "Primary Team"],
    [
        ["01", "Org Tree and Entity Hierarchy",        "Weeks 1–2",   "Backend + DB"],
        ["02", "RBAC Engine and Permission Matrix",    "Weeks 3–4",   "Backend + Security"],
        ["03", "Member Onboarding Wizard",             "Weeks 5–6",   "Full-Stack"],
        ["04", "Work OS Data Models",                  "Weeks 7–8",   "Backend + DB"],
        ["05", "Task CRUD and List View",              "Weeks 9–10",  "Full-Stack"],
        ["06", "Board View, Dependencies, Real-Time",  "Weeks 11–12", "Full-Stack + Infra"],
        ["07", "Employee Profile and Directory",       "Weeks 13–14", "Full-Stack"],
        ["08", "Leave Management Engine",              "Weeks 15–16", "Full-Stack"],
        ["09", "Attendance, Holidays, Team Calendar",  "Weeks 17–18", "Full-Stack"],
    ],
    col_widths=[0.6, 3.0, 1.2, 2.3]
)

# ── SPRINT 01 ────────────────────────────────────────────────────────────
sprint_header("A", 1, "Org Tree and Entity Hierarchy", "Weeks 1–2",
"Design, implement, and test the full Organisation and Entity Hierarchy data model. By the end of this sprint, the system must support the creation of an Organisation (the top-level tenant), Departments and Sub-Departments in a recursive tree, Teams within Departments, and the assignment of Users to Teams with a Role. Every subsequent module depends on this hierarchy, so the data model must be designed for extensibility from day one.")

h2("8.1 Sprint 01 — Org Tree and Entity Hierarchy")
body("The organisational hierarchy is the skeleton of Slipwise. Every other piece of data — tasks, deals, tickets, documents — is ultimately owned by either an Organisation or a Member of an Organisation. Getting this model right is the most important architectural decision in Phase A. The hierarchy has four levels: Organisation, Department, Sub-Department (optional), and Team. A User can belong to multiple Teams across different Departments, but they have exactly one primary Department for HR reporting purposes.")

sprint_story(1, "Create Organisation Record",
    "a Slipwise platform administrator",
    "create a new Organisation with a name, slug, logo, timezone, and primary currency",
    "a new tenant is provisioned with its own isolated data space",
    5, "P0 — Must Have")
sprint_acceptance([
    "POST /api/v1/organisations creates an Organisation record with all required fields.",
    "The slug must be unique across all organisations and contain only lowercase alphanumeric characters and hyphens.",
    "The timezone field is validated against the IANA timezone database.",
    "The currency field is validated against the ISO 4217 currency code list.",
    "The creating user is automatically assigned the OWNER role within the new organisation.",
    "A default 'General' Department is automatically created for the organisation.",
    "The operation is wrapped in a database transaction — either the Organisation and default Department both exist, or neither does.",
])
sprint_tech([
    "Use Prisma's createMany inside a $transaction to atomically create Organisation + default Department.",
    "Slug uniqueness enforced at DB level with a unique index on Organisation.slug.",
    "Timezone validation uses the Intl.supportedValuesOf('timeZone') API in a Zod refinement.",
])

sprint_story(2, "Create and Manage Departments",
    "an Organisation administrator",
    "create top-level Departments and optional Sub-Departments forming a tree",
    "the company's reporting structure is accurately reflected in the system",
    8, "P0 — Must Have")
sprint_acceptance([
    "POST /api/v1/departments creates a Department optionally under a parent Department.",
    "The parentId field is optional; null parentId denotes a top-level Department.",
    "Departments cannot be nested more than 4 levels deep (enforced at API layer).",
    "GET /api/v1/departments returns the full tree structure in a single response using a recursive CTE query.",
    "Deleting a Department that contains active members returns a 409 Conflict error with a list of affected members.",
    "Moving a Department to a new parent checks for circular reference and returns 400 if detected.",
    "Each Department can have a designated Head (a UserId) shown in the org chart.",
])
sprint_tech([
    "Use PostgreSQL recursive Common Table Expression (WITH RECURSIVE) for fetching the full tree in one DB round-trip.",
    "Store parent_department_id as a nullable foreign key on the Department table.",
    "Circular reference check: before setting a new parentId, verify the proposed new parent is not a descendant of the current node.",
])

sprint_story(3, "Create Teams within Departments",
    "a Department manager",
    "create named Teams within my Department and assign members to them",
    "work can be assigned and reported at the Team level",
    5, "P0 — Must Have")
sprint_acceptance([
    "POST /api/v1/teams creates a Team linked to a Department.",
    "A Team must have a name, an optional description, and a Team Lead (UserId).",
    "A User can be a member of multiple Teams across the organisation.",
    "GET /api/v1/departments/:id/teams returns all teams for a department with their member counts.",
    "POST /api/v1/teams/:id/members adds a User to a Team.",
    "DELETE /api/v1/teams/:id/members/:userId removes a User from a Team.",
    "Removing the Team Lead triggers a warning if no replacement Lead is specified.",
])
sprint_tech([
    "TeamMembership is a join table: (teamId, userId, joinedAt). No duplicates enforced by composite unique index.",
    "Team Lead is stored as a teamLeadUserId nullable column on the Team table, not a separate join.",
])

sprint_story(4, "Org Chart Visual View",
    "any organisation member with the 'view_org_chart' permission",
    "view an interactive, collapsible org chart of the organisation",
    "I can quickly understand reporting lines and find who is in which department",
    8, "P1 — Should Have")
sprint_acceptance([
    "The org chart page is accessible at /app/people/org-chart.",
    "The chart renders the full Department tree with Department names, heads, and member counts.",
    "Clicking a Department node expands it to show Teams and their members.",
    "Clicking a person's avatar opens a side panel with their profile summary.",
    "The chart is searchable — typing a name highlights matching nodes and scrolls to them.",
    "The chart can be exported as a PNG image via a browser-based canvas render.",
    "Mobile view renders the org chart as a collapsible list, not a horizontal tree.",
])
sprint_tech([
    "Use D3.js tree layout for the org chart rendering. Data is fetched from GET /api/v1/org-chart which returns the pre-computed tree JSON.",
    "Server-side cache the org chart JSON in Redis with a 5-minute TTL — invalidate on any Department/Team/Membership change.",
])
page_break()

# ── SPRINT 02 ────────────────────────────────────────────────────────────
sprint_header("A", 2, "RBAC Engine and Permission Matrix", "Weeks 3–4",
"Build the complete Role-Based Access Control (RBAC) engine including the role definition system, the granular permission grant table, the permission check middleware, and the administrator UI for the Permission Matrix. By end of sprint, an administrator must be able to open the Permission Matrix panel, see every permission in the system, and grant or revoke them per user.")

h2("8.2 Sprint 02 — RBAC Engine and Permission Matrix")
body("The RBAC system is the single most security-critical component in Slipwise. Every feature, every page, every API endpoint depends on it to decide what a user is allowed to do. The system must be both granular enough to support fine-grained control and performant enough that it does not add perceptible latency to every request. The architecture uses a three-layer model: System Roles (Owner, Admin, Member, Guest) define a base set of permissions, Role Templates allow administrators to define named permission bundles, and individual GranularPermissionGrant records override the template for a specific user.")

sprint_story(5, "Define System Roles",
    "the Slipwise platform",
    "enforce four built-in system roles — Owner, Admin, Member, and Guest — with predefined base permissions",
    "there is always a safe default permission level for new users",
    5, "P0 — Must Have")
sprint_acceptance([
    "The four system roles are seeded in the database during migration and cannot be deleted.",
    "OWNER: has all permissions. Only one Owner per Organisation. The Owner role cannot be removed from the user who created the Organisation.",
    "ADMIN: has all permissions except billing management and deleting the Organisation.",
    "MEMBER: has read access to their Department's content and write access to tasks assigned to them.",
    "GUEST: has read-only access to specifically shared content. Cannot create any records.",
    "Assigning a user to a higher role does not require the RBAC engine to be re-run — role is checked at request time.",
])
sprint_tech([
    "System roles are stored in a SystemRole enum and referenced by the OrganisationMember.systemRole field.",
    "Permission check function resolves: GranularPermissionGrant (user-specific) → RoleTemplate grants → SystemRole base permissions. First matching record wins.",
])

sprint_story(6, "GranularPermissionGrant Table and API",
    "an organisation administrator",
    "explicitly grant or revoke specific permissions for individual users",
    "I have fine-grained control over what each person can access beyond their system role",
    13, "P0 — Must Have")
sprint_acceptance([
    "The GranularPermissionGrant table has columns: id, organisationId, userId, permissionKey, granted (boolean), grantedBy, grantedAt.",
    "POST /api/v1/permissions/grant creates or updates a grant record.",
    "A grant with 'granted: false' is an explicit deny, which overrides any role-based permission.",
    "GET /api/v1/permissions/users/:userId returns all grants for a user in the requesting organisation.",
    "Permissions are cached in Redis per (userId, organisationId) with a 60-second TTL. Cache is invalidated immediately on any grant change.",
    "The permission key namespace is documented (see Appendix): e.g., 'mailbox:read', 'invoices:create', 'crm:deals:view'.",
    "Every permission key that can be granted is defined in a TypeScript permissions.ts constants file — no magic strings.",
])
sprint_tech([
    "The checkPermission(userId, orgId, key) function: (1) check Redis cache, (2) if miss, query GranularPermissionGrant for explicit grant/deny, (3) if no explicit record, check user's RoleTemplate, (4) if no template, apply SystemRole defaults. Cache the resolved result.",
    "Composite unique index on (organisationId, userId, permissionKey) for upsert performance.",
])

sprint_story(7, "Permission Matrix UI — Administrator View",
    "an organisation administrator",
    "open a Permission Matrix UI panel when adding or editing a user and see every permission listed with toggle controls",
    "I can easily control what each user has access to without needing to know internal permission key names",
    13, "P1 — Should Have")
sprint_acceptance([
    "The Permission Matrix panel is accessible from the People module when viewing or editing a user.",
    "Permissions are grouped by module: Work OS, Mailbox, Docs & Invoices, CRM, HRIS, ITSM, Reports, Admin.",
    "Within each module, sub-permissions are listed (e.g., within CRM: View Contacts, Create Contacts, Edit Contacts, Delete Contacts, View Deals, Manage Pipelines, etc.).",
    "Each permission row shows a toggle switch (On/Off) and an indicator showing whether the permission comes from the user's system role, role template, or an individual override.",
    "Toggling a permission creates or updates a GranularPermissionGrant record via the API.",
    "Changes take effect within 60 seconds (the Redis cache TTL).",
    "A 'Reset to Role Defaults' button removes all individual overrides for the user.",
    "For Mailbox permissions: the panel shows a list of all organisation mailboxes and the user can toggle access per mailbox.",
])
sprint_tech([
    "Frontend: virtualize the permission list with react-virtual if >100 rows to prevent DOM bloat.",
    "Use optimistic updates — toggle flips immediately in UI, API call in background. Revert on error.",
])

sprint_story(8, "Role Templates",
    "an organisation administrator",
    "create named Role Templates (e.g., 'Sales Rep', 'Support Agent', 'Finance Viewer') that bundle a set of permissions",
    "I can assign a template to multiple users instead of configuring permissions individually each time",
    8, "P1 — Should Have")
sprint_acceptance([
    "POST /api/v1/role-templates creates a named template with a set of permissionKey/granted pairs.",
    "Role templates are organisation-scoped — not shared across tenants.",
    "Assigning a template to a user applies all the template's grants as defaults.",
    "User-specific GranularPermissionGrant records override the template.",
    "Editing a template propagates permission changes to all users assigned to that template (async job).",
    "A template cannot be deleted if users are currently assigned to it (returns 409 with user count).",
    "Three default templates are seeded per organisation: 'Standard Member', 'Read Only', and 'Department Manager'.",
])
page_break()

# ── SPRINT 03 ────────────────────────────────────────────────────────────
sprint_header("A", 3, "Member Onboarding Wizard", "Weeks 5–6",
"Build the complete Member Onboarding Wizard flow that allows administrators to invite new team members, assign them to departments and teams, set their role and role template, configure their permissions via the Permission Matrix, and send them a personalised invitation email. The wizard must be polished, fast, and reduce the time to add a new team member to under 5 minutes.")

h2("8.3 Sprint 03 — Member Onboarding Wizard")
body("When a new employee joins a company, setting them up in the system must be fast and error-free. The onboarding wizard is a multi-step, guided process that walks administrators through every decision — from the basic invite details to the exact set of mailboxes the new hire should access. The wizard is also the primary way that permission complexity is exposed to administrators in a friendly, contextual way.")

sprint_story(9, "Invite Member via Email",
    "an organisation administrator",
    "enter a new team member's email address and send them an invitation",
    "they can accept it and create their Slipwise account linked to our organisation",
    5, "P0 — Must Have")
sprint_acceptance([
    "POST /api/v1/members/invite accepts { email, firstName, lastName, roleTemplateId } and creates a pending Invitation record.",
    "An invitation email is sent via AWS SES with a unique, time-limited (72 hours) magic link.",
    "If the email is already associated with a Slipwise account, the invitation links the existing account to the new organisation.",
    "Re-inviting a pending invitee resends the email with a fresh token and extends the expiry.",
    "The invitation list is visible on the People module with status: Pending, Accepted, Expired.",
    "Expired invitations can be resent or revoked.",
])
sprint_tech([
    "Invitation tokens are cryptographically random (32 bytes, hex-encoded). Stored as a hashed value in the Invitation table. The plaintext token is sent in the email link only.",
    "Invitation acceptance flow: GET /invite/:token → validate → create User + OrganisationMember → redirect to onboarding profile setup.",
])

sprint_story(10, "Wizard Step 1 — Basic Details",
    "an organisation administrator using the onboarding wizard",
    "enter the new member's name, personal email, work email, phone number, job title, employment type, and start date",
    "the member's basic HR record is created correctly from day one",
    5, "P0 — Must Have")
sprint_acceptance([
    "Step 1 of the wizard shows a clean form with: First Name, Last Name, Personal Email, Work Email, Phone Number, Job Title, Employment Type (Full-time / Part-time / Contractor / Intern), Start Date.",
    "All required fields are validated inline before the user can proceed to step 2.",
    "The wizard auto-saves a draft on every step change so the admin can return and resume.",
    "Email format validation is performed on both email fields.",
    "Start date cannot be more than 90 days in the past.",
])

sprint_story(11, "Wizard Step 2 — Org Assignment",
    "an organisation administrator using the onboarding wizard",
    "assign the new member to a Department, an optional Sub-Department, and one or more Teams",
    "the member appears in the correct place in the org chart and team rosters from day one",
    5, "P0 — Must Have")
sprint_acceptance([
    "Step 2 shows a searchable dropdown for Department (required) and Sub-Department (optional, filtered by selected Department).",
    "A multi-select field allows adding the member to one or more Teams within the selected Department.",
    "The Team Lead toggle is available per team if the current Team Lead slot is empty.",
    "The step shows a live preview of how the member will appear in the org chart.",
])

sprint_story(12, "Wizard Step 3 — Permission Setup",
    "an organisation administrator using the onboarding wizard",
    "select a Role Template for the new member and then optionally customise individual permissions using the Permission Matrix panel",
    "the member has exactly the right access from their first login",
    8, "P1 — Should Have")
sprint_acceptance([
    "Step 3 shows the Role Template selector (defaults to 'Standard Member').",
    "Below the template selector, the full Permission Matrix panel is embedded, pre-filled based on the selected template.",
    "The admin can expand each module section and toggle individual permissions.",
    "Mailbox access configuration is shown as a sub-section under the Permission Matrix: a list of all org mailboxes with toggle per mailbox.",
    "A summary panel on the right side shows a human-readable list of 'What this user CAN do' and 'What this user CANNOT do'.",
    "Permissions are not saved until the wizard is completed — they are held in the wizard's session state.",
])

sprint_story(13, "Wizard Step 4 — Review and Send Invite",
    "an organisation administrator using the onboarding wizard",
    "review a complete summary of the new member's details, org assignment, and permissions before finalising",
    "I can catch mistakes before they cause issues for the new hire",
    5, "P0 — Must Have")
sprint_acceptance([
    "Step 4 shows a structured summary: Personal Details, Org Assignment, System Role, Role Template, and Permission Overrides.",
    "Each section has an 'Edit' link that returns to that step.",
    "A 'Send Invitation' button triggers: (1) create pending member record, (2) apply permissions, (3) send invitation email.",
    "After sending, the admin sees a confirmation screen with the invite link for manual sharing if needed.",
    "The invitation can be cancelled from the confirmation screen for 5 minutes (grace period).",
])
page_break()

# ── SPRINT 04 ────────────────────────────────────────────────────────────
sprint_header("A", 4, "Work OS Data Models", "Weeks 7–8",
"Design and implement the complete Work OS data model: Workspace configuration, Space, Folder, List, Task, Subtask, Comment, Attachment, Custom Status, Custom Field definition, and Task Custom Field Value. All models must be created with their CRUD API routes. No UI is required this sprint — the focus is a rock-solid, extensible data foundation.")

h2("8.4 Sprint 04 — Work OS Data Models")
body("The Work OS is Slipwise's answer to ClickUp and Asana. To compete with these best-in-class tools, the data model must be designed from the start to support every view (List, Board, Gantt, Calendar, Workload, Table, Form, Mind Map, Timeline, Dashboard) and every customisation feature (custom statuses, custom fields, recurring tasks, task dependencies, time estimates) without requiring schema migrations later. This sprint invests in getting the foundation right.")

sprint_story(14, "Space Model — CRUD",
    "an organisation administrator or member with Space creation permission",
    "create, read, update, and archive Spaces representing departments or major project areas",
    "the organisation's work is organised at the highest level into logical containers",
    8, "P0 — Must Have")
sprint_acceptance([
    "A Space has: id, organisationId, name, description, color, icon (emoji or icon name), isPrivate, createdBy, archivedAt.",
    "Private Spaces are only visible to their explicit members (SpaceMember join table).",
    "POST /api/v1/spaces creates a Space and automatically adds the creator as a Space Admin.",
    "GET /api/v1/spaces returns all Spaces visible to the requesting user (public + their private spaces).",
    "PATCH /api/v1/spaces/:id updates name, description, color, icon, or privacy setting.",
    "DELETE /api/v1/spaces/:id is a soft delete (sets archivedAt). Archived spaces are hidden by default but accessible via a filter.",
    "On Space creation, three default Custom Statuses are created: 'To Do', 'In Progress', 'Done'.",
])
sprint_tech([
    "isPrivate=true: API enforces SpaceMember check. Public spaces: all org members can see.",
    "Space archival is a soft delete. Add archivedAt index for fast filtering of active spaces.",
])

sprint_story(15, "List Model — CRUD",
    "a Space member",
    "create Lists inside a Space or optionally inside a Folder to organise tasks into discrete projects or queues",
    "tasks are organised into meaningful collections that match how my team works",
    8, "P0 — Must Have")
sprint_acceptance([
    "A List has: id, spaceId, folderId (nullable), name, description, color, defaultStatus, sortOrder.",
    "POST /api/v1/lists creates a List inside a Space or Folder.",
    "Lists can be reordered within their parent (Space or Folder) by updating sortOrder.",
    "GET /api/v1/spaces/:id/lists returns all Lists in a Space, grouped by Folder.",
    "Archiving a List soft-deletes it and all its Tasks.",
    "A List can be cloned: POST /api/v1/lists/:id/clone creates a new List with the same configuration and optionally copies tasks.",
])

sprint_story(16, "Task Model — Full Schema",
    "the engineering team",
    "have a Task data model that supports all task attributes needed for all planned Work OS views",
    "no schema migrations are needed when new views are introduced in future sprints",
    13, "P0 — Must Have")
sprint_acceptance([
    "A Task has: id, listId, parentTaskId (nullable, for subtasks), title, description (rich text JSON), assigneeId, reporterId, statusId, priority (NONE/LOW/MEDIUM/HIGH/URGENT), dueDate, startDate, timeEstimate (minutes), timeTracked (minutes), sortOrder, tags[], isRecurring, recurringConfig (JSON).",
    "POST /api/v1/tasks creates a Task. The statusId must belong to the Task's Space.",
    "Tasks with parentTaskId set are Subtasks — they count toward the parent's completion percentage.",
    "Parent task completion percentage = (completed subtasks / total subtasks) × 100.",
    "Task tags are stored as a PostgreSQL text[] array with a GIN index for fast tag filtering.",
    "Recurring task config stores: frequency (DAILY/WEEKLY/MONTHLY), interval, nextDueDate, endDate.",
])
sprint_tech([
    "Task description uses the ProseMirror JSON format (compatible with Tiptap editor).",
    "Add a tsvector column on Task (title + description) with a GIN index for full-text search.",
    "TaskDependency is a separate join table: (predecessorTaskId, successorTaskId, dependencyType: FINISH_TO_START / START_TO_START / FINISH_TO_FINISH).",
])

sprint_story(17, "Custom Statuses and Custom Fields",
    "a Space administrator",
    "define custom statuses and custom fields for my Space",
    "tasks in my Space can have the exact metadata fields my team needs",
    13, "P1 — Should Have")
sprint_acceptance([
    "Custom Status: id, spaceId, name, color (hex), type (OPEN/IN_PROGRESS/CLOSED/CUSTOM), sortOrder.",
    "POST /api/v1/spaces/:id/statuses creates a custom status.",
    "Custom Field Definition: id, spaceId, name, type (TEXT/NUMBER/DATE/DROPDOWN/PERSON/URL/CHECKBOX/RATING/CURRENCY/FORMULA), config (JSON, e.g. dropdown options).",
    "TaskCustomFieldValue: id, taskId, fieldDefinitionId, value (stored as JSON, type-specific).",
    "GET /api/v1/spaces/:id/fields returns all custom field definitions for a space.",
    "Custom fields are shown in the task detail panel and in Table view as configurable columns.",
])
page_break()

# ── SPRINT 05 ────────────────────────────────────────────────────────────
sprint_header("A", 5, "Task CRUD and List View", "Weeks 9–10",
"Build the complete List View for the Work OS — the primary way users interact with tasks. This includes the task creation flow, task detail panel, inline editing, filtering, sorting, grouping, and the complete API layer for all task operations.")

h2("8.5 Sprint 05 — Task CRUD and List View")
body("The List View is the most-used view in any project management tool. It must be fast to render even with hundreds of tasks, support keyboard shortcuts, allow drag-and-drop reordering, and provide a rich task detail panel without full page navigation. The performance requirement is that the List View must render the first visible viewport of tasks within 200ms of navigation.")

sprint_story(18, "Task List View — Render and Pagination",
    "a Work OS user",
    "view all tasks in a List in a clear, fast-loading list layout with task titles, status, priority, assignee, and due date",
    "I can quickly scan my team's work without needing to open individual tasks",
    13, "P0 — Must Have")
sprint_acceptance([
    "The List View is accessible at /app/work-os/lists/:listId.",
    "Tasks are paginated: 50 tasks per page, with infinite scroll loading the next batch.",
    "Each row shows: checkbox (for bulk actions), task title (clickable to open detail panel), custom status badge (coloured), priority icon, assignee avatar, due date (red if overdue), tags.",
    "Tasks can be grouped by: Status, Assignee, Priority, Due Date, or Custom Dropdown Field. Grouping collapses/expands each group.",
    "Sorting is available on: Due Date, Priority, Created Date, Alphabetical. Sorting and grouping can be combined.",
    "A filter bar allows filtering by: Assignee (multi-select), Status (multi-select), Priority (multi-select), Due Date range, Tag, Custom Field value.",
    "Active filters are shown as dismissible chips below the filter bar.",
    "Filter state is persisted in the URL query string so it can be shared via link.",
])
sprint_tech([
    "Use cursor-based pagination (after: lastTaskId) rather than LIMIT/OFFSET for consistent results under concurrent inserts.",
    "Filters are translated to a Prisma where clause. Use a query builder utility function to map filter chips to Prisma conditions.",
])

sprint_story(19, "Task Detail Panel",
    "a Work OS user",
    "open a task detail panel (slide-in drawer) when clicking a task without losing my place in the List View",
    "I can read and edit full task details without navigating away",
    13, "P0 — Must Have")
sprint_acceptance([
    "Clicking a task title in the List View opens a slide-in right panel without a page reload.",
    "The panel shows: Task title (editable inline), Description (rich text editor, Tiptap), Status (dropdown), Priority (dropdown), Assignee (user picker), Reporter, Due Date (date picker), Start Date, Time Estimate, Time Tracked, Tags (multi-input), Subtask list, Attachment list, Comment thread, Activity log.",
    "All fields in the panel are editable inline with auto-save on blur or Enter.",
    "Subtasks can be created, checked off, and assigned directly from the panel.",
    "The comment thread supports @mentions (resolves to org members), file attachments, and emoji reactions.",
    "Activity log shows: field changes (who changed what field from what value to what value), comments, attachments, and status transitions with timestamps.",
    "The panel is deep-linkable: /app/work-os/tasks/:taskId opens the parent List view with the panel pre-opened.",
])
sprint_tech([
    "Rich text (Tiptap): store as ProseMirror JSON. Render to HTML on the fly — never store HTML directly.",
    "Activity log: on every PATCH to a task, compute a diff of changed fields and write AuditLog records. Fetch activity log separately from the task detail.",
])

sprint_story(20, "Task Create — Quick Create and Full Create",
    "a Work OS user",
    "create a task either via a quick one-line input or a full creation modal with all fields",
    "I can capture work quickly or add full details depending on my situation",
    8, "P0 — Must Have")
sprint_acceptance([
    "Quick Create: a '+' button or pressing 'T' on the keyboard opens a single-line input at the bottom of the current group. Pressing Enter creates the task and opens another quick-create input.",
    "Full Create: a modal with all task fields, accessed via the '+ New Task' button or by pressing Shift+T.",
    "Both flows pre-fill the List, Status (first status in Space), and Reporter (current user).",
    "The task is optimistically added to the list immediately, before the API call returns.",
    "On API failure, the optimistically added task is removed and an error toast is shown.",
])

sprint_story(21, "Bulk Task Operations",
    "a Work OS user",
    "select multiple tasks using checkboxes and perform bulk operations on them",
    "I can efficiently update many tasks at once rather than editing them one by one",
    8, "P1 — Should Have")
sprint_acceptance([
    "Selecting any task's checkbox shows a bulk action toolbar at the bottom of the screen.",
    "Bulk actions available: Change Status, Change Assignee, Change Priority, Add Tag, Set Due Date, Move to another List, Delete.",
    "A 'Select All' checkbox in the header selects all tasks visible on the current page.",
    "Bulk delete requires a confirmation dialog showing the number of tasks to be deleted.",
    "Bulk actions are processed server-side in a single batch API call (not one call per task).",
])
page_break()

# ── SPRINT 06 ────────────────────────────────────────────────────────────
sprint_header("A", 6, "Board View, Task Dependencies, and Real-Time Updates", "Weeks 11–12",
"Build the Kanban Board View, implement task dependencies with visual blocking indicators, and introduce real-time collaborative updates using Supabase Realtime and Redis Pub/Sub so that changes made by one user are reflected on other users' screens within 2 seconds.")

h2("8.6 Sprint 06 — Board View, Dependencies, Real-Time")
body("The Board (Kanban) View is the second most commonly used task view. It shows tasks as cards arranged in columns, where each column represents a status. Users can drag cards between columns to change status, and drag within a column to change sort order. Real-time updates are critical here — when one team member moves a card, all other team members viewing the same board must see the change without refreshing.")

sprint_story(22, "Kanban Board View",
    "a Work OS user",
    "switch to a Board View that shows tasks as cards in columns representing each status",
    "I can see the flow of work across statuses at a glance and move tasks by dragging",
    13, "P0 — Must Have")
sprint_acceptance([
    "The Board View is accessible by clicking the 'Board' tab in the view selector at the top of a List.",
    "Each status column shows: column header with status name, color badge, and task count. Tasks within the column as cards.",
    "Task cards show: title, priority icon, assignee avatar, due date, tag chips, subtask progress bar (e.g., '2/5 subtasks done').",
    "Cards can be dragged between columns (changes status) and within a column (changes sort order). Both operations call the API on drop.",
    "A card can be clicked to open the Task Detail Panel.",
    "Columns can be collapsed (clicking the column header) to save horizontal space.",
    "A '+' button at the bottom of each column opens a quick-create input pre-filled with that column's status.",
    "The Board View respects the same filters as the List View — filtered tasks are hidden from the Board.",
])
sprint_tech([
    "Use @hello-pangea/dnd (DnD kit successor) for drag-and-drop. Handle reordering via sortOrder floating-point algorithm (insert between two numbers without full re-index).",
    "On drop: optimistic update immediately, PATCH /api/v1/tasks/:id with new statusId and sortOrder in background.",
])

sprint_story(23, "Task Dependencies",
    "a Work OS user",
    "link tasks with dependency relationships (Finish-to-Start, Start-to-Start, Finish-to-Finish)",
    "the system indicates when a task cannot be started because its predecessor is not yet complete",
    8, "P1 — Should Have")
sprint_acceptance([
    "From the Task Detail Panel, a 'Dependencies' section allows adding predecessor or successor tasks.",
    "Dependency types supported: Finish-to-Start (default), Start-to-Start, Finish-to-Finish.",
    "When a task has an incomplete predecessor (Finish-to-Start dependency), the task card shows a 'Blocked' red badge.",
    "Blocked tasks cannot be moved to 'In Progress' or later statuses (configurable per Space).",
    "The Gantt View (built in Sprint 12) renders dependencies as connecting lines.",
    "Circular dependency detection: the API returns a 400 error if adding a dependency would create a cycle.",
])
sprint_tech([
    "TaskDependency table: predecessorTaskId, successorTaskId, type. Composite unique index prevents duplicates.",
    "Circular dependency check: BFS/DFS traversal of the dependency graph before inserting. Maximum depth limit of 20 to prevent performance issues.",
])

sprint_story(24, "Real-Time Task Updates",
    "a Work OS user",
    "see task status changes, new tasks, and comment additions made by other users reflected on my screen within 2 seconds without refreshing",
    "my team can collaborate in real-time without confusion about stale data",
    13, "P0 — Must Have")
sprint_acceptance([
    "When User A changes a task's status while User B is viewing the same Board, User B sees the card move to the new column within 2 seconds.",
    "When a new task is created in a List, all users viewing that List see the new task appear at the correct position.",
    "When a comment is added to a task, the task detail panel on all other viewers shows the new comment immediately.",
    "The system uses Supabase Realtime channels named by List ID (e.g., 'list:clm_abc123').",
    "Each client subscribes to the relevant channels on mount and unsubscribes on unmount.",
    "A connection indicator in the bottom-left shows: green dot (connected), yellow dot (reconnecting), red dot (disconnected, changes may not be live).",
])
sprint_tech([
    "Architecture: Task mutations trigger Supabase Realtime broadcasts via a Postgres trigger on the Task table (INSERT, UPDATE, DELETE).",
    "For high-frequency events (comment floods), use Redis Pub/Sub to debounce and batch updates.",
    "Client: use Supabase's useChannel hook. On 'task:updated' event, update the local SWR/React Query cache directly without a refetch.",
])
page_break()

# ── SPRINT 07 ────────────────────────────────────────────────────────────
sprint_header("A", 7, "Employee Profile and Directory", "Weeks 13–14",
"Build the Employee Profile system — the rich, comprehensive profile page for each team member — and the People Directory with search, filters, and export capabilities. By end of sprint, every employee can view their own profile and browse the organisation directory.")

h2("8.7 Sprint 07 — Employee Profile and Directory")
body("The Employee Profile is the source of truth for everything about a person in Slipwise. It goes far beyond a simple name and email — it captures work history within the company, documents, emergency contacts, equipment assigned, performance review history, and current assignments. The directory lets anyone find anyone in the company quickly.")

sprint_story(25, "Employee Profile — Core Fields",
    "any organisation member",
    "view a comprehensive profile page for any team member showing their personal, professional, and organisational information",
    "I always have one place to go to find any information about a colleague",
    13, "P0 — Must Have")
sprint_acceptance([
    "Profile page at /app/people/:memberId covers: Photo, Name, Job Title, Department, Teams, System Role, Employment Type, Start Date, Work Email, Personal Email, Phone, Time Zone, Pronouns (optional), Short Bio.",
    "A 'Work' section shows: Manager (direct report chain), Direct Reports, Current Task Assignments (top 5 open tasks), Active Spaces.",
    "An 'Employment' section shows: Employment History within the org (title changes, promotions), Contract Type, Salary Band (visible only to HR/Admin), Notice Period.",
    "A 'Documents' section shows: uploaded documents (offer letter, ID, certifications). Access controlled by permission.",
    "An 'Emergency Contacts' section: Name, Relationship, Phone. Visible only to HR/Admin.",
    "An 'Equipment' section: list of company equipment assigned (device type, serial number, assigned date).",
    "The profile has a public vs. private field model — employees control which fields are visible to all members vs. HR only.",
])
sprint_tech([
    "EmployeeProfile is a 1:1 extension table from the User table. Keeps the User table clean and the profile table HR-specific.",
    "Sensitive fields (salary, ID docs, emergency contacts) have a separate permission check beyond the base 'view_employee_profile' permission.",
])

sprint_story(26, "People Directory",
    "any organisation member",
    "search the company directory to find colleagues by name, department, team, or job title",
    "I can quickly find and contact the right person without asking around",
    8, "P0 — Must Have")
sprint_acceptance([
    "The directory at /app/people is a grid of employee cards with photo, name, title, department.",
    "A search bar filters in real-time (client-side on loaded data, server-side search for >500 employees).",
    "Filter options: Department (multi-select), Team (multi-select), Employment Type (multi-select), Location (if configured).",
    "Clicking a card opens the Employee Profile page.",
    "An export button generates a CSV of the filtered directory (name, email, title, department). Requires 'export_directory' permission.",
    "The directory shows a count of total active employees.",
])
page_break()

# ── SPRINT 08 ────────────────────────────────────────────────────────────
sprint_header("A", 8, "Leave Management Engine", "Weeks 15–16",
"Build the complete Leave Management system: leave type configuration, leave balance calculation and accrual, the leave request workflow with manager approval, leave calendar, and all associated notifications.")

h2("8.8 Sprint 08 — Leave Management Engine")
body("Leave management is one of the most process-heavy and compliance-sensitive features in any HRIS. Errors in leave balance calculation directly impact payroll and employee trust. The system must handle different leave types (annual, sick, maternity, paternity, compensatory, unpaid), different accrual models (fixed allowance, monthly accrual, anniversary-based), carry-over rules, and a complete approval workflow.")

sprint_story(27, "Leave Type Configuration",
    "an HR administrator",
    "configure leave types for my organisation including accrual rules, carry-over limits, and eligibility criteria",
    "the system automatically calculates correct balances for each employee type",
    8, "P0 — Must Have")
sprint_acceptance([
    "POST /api/v1/leave-types creates a leave type with: name, code, description, paidOrUnpaid, accrualType (FIXED_ANNUAL/MONTHLY_ACCRUAL/ANNIVERSARY), annualAllowance (days), carryOverLimit (days, 0 = no carry over), maxCarryOverExpiryDate, eligibleEmploymentTypes[].",
    "Pre-seeded leave types on org creation: Annual Leave (20 days, FIXED_ANNUAL), Sick Leave (10 days), Maternity Leave (90 days), Paternity Leave (15 days), Unpaid Leave (unlimited).",
    "MONTHLY_ACCRUAL type calculates monthly accrual as annualAllowance / 12 added to balance on the 1st of each month (BullMQ cron job).",
    "ANNIVERSARY type adds the full annual allowance on the employee's work anniversary date.",
    "Carry-over balance is calculated and posted on the organisation's new financial year start date.",
])
sprint_tech([
    "LeaveBalance table: (employeeId, leaveTypeId, year, totalAllowance, accrued, used, carryOver, available). Computed nightly by a BullMQ cron job for accrual types.",
    "For FIXED_ANNUAL, the balance is simply set at the start of the year based on the allowance.",
])

sprint_story(28, "Leave Request Workflow",
    "an employee",
    "submit a leave request specifying the leave type, start date, end date, and an optional note",
    "my manager is notified and can approve or reject my request",
    13, "P0 — Must Have")
sprint_acceptance([
    "POST /api/v1/leave-requests creates a request in PENDING status.",
    "The system calculates the number of working days (excluding weekends and public holidays) in the requested range.",
    "If the calculated days exceed the employee's available balance, the API returns a warning (not a block — some orgs allow negative balance).",
    "The employee's direct manager (and any configured backup approver) receives an email + in-app notification immediately.",
    "The manager can APPROVE or REJECT with a mandatory reason from the notification email's action buttons or from the Leave module UI.",
    "On APPROVE: the employee's LeaveBalance.used is incremented and they receive an approval notification.",
    "On REJECT: the employee receives a rejection notification with the reason.",
    "An employee can CANCEL a pending or approved request. Cancelling an approved request automatically credits back the used days.",
    "Leave requests cannot overlap with existing approved requests for the same employee.",
])
sprint_tech([
    "Approval workflow state machine: PENDING → APPROVED / REJECTED / CANCELLED. No direct jumps between non-adjacent states.",
    "Working day calculation: exclude weekends + records in the HolidayCalendar table for the organisation's country.",
])

sprint_story(29, "Leave Balance Dashboard",
    "an employee",
    "view my current leave balances for all leave types and a history of my past requests",
    "I always know exactly how much leave I have available before submitting a request",
    8, "P0 — Must Have")
sprint_acceptance([
    "The Leave dashboard at /app/hris/leave shows: a card per leave type with total allowance, used, carried over, and available days shown as a donut chart.",
    "A 'Pending Requests' section shows all my submitted requests awaiting approval.",
    "A 'History' table shows the last 24 months of approved leave with date ranges and duration.",
    "HR admins see a team-wide leave overview: who is on leave today, who has requests pending approval.",
    "A 'Team Calendar' tab shows all approved team leave in a calendar view (built further in Sprint 09).",
])
page_break()

# ── SPRINT 09 ────────────────────────────────────────────────────────────
sprint_header("A", 9, "Attendance, Holiday Calendar, and Team Calendar", "Weeks 17–18",
"Build attendance tracking with clock-in/clock-out, the organisation holiday calendar with country-based presets, the team leave calendar, and the first version of the global Team Calendar that shows leave, work anniversaries, and key HR dates.")

h2("8.9 Sprint 09 — Attendance, Holidays, Team Calendar")
body("Attendance tracking provides visibility into daily presence and working hours. Combined with the Holiday Calendar and the Team Calendar, it gives HR teams and managers a complete view of workforce availability on any given day. This sprint completes Phase A.")

sprint_story(30, "Attendance Clock-In / Clock-Out",
    "an employee",
    "clock in when I start work and clock out when I finish from within Slipwise",
    "my attendance is automatically recorded and my daily hours are tracked",
    8, "P0 — Must Have")
sprint_acceptance([
    "A persistent widget in the Slipwise sidebar shows the current clock status: 'Clocked In' (green, time elapsed) or 'Clocked Out' (grey, last clock-out time).",
    "POST /api/v1/attendance/clock-in creates an AttendanceRecord with clockInTime = now().",
    "POST /api/v1/attendance/clock-out updates the open record with clockOutTime = now() and calculates totalMinutes.",
    "An employee can only have one open attendance record at a time (clock-in without a clock-out).",
    "If an employee forgets to clock out, the system automatically closes the record at midnight and flags it as 'auto-closed'.",
    "Manual attendance correction by HR: HR can edit clock-in/out times with a mandatory reason note.",
    "Attendance records are exported as a CSV per employee or for the whole team per month.",
])

sprint_story(31, "Holiday Calendar",
    "an HR administrator",
    "configure the organisation's holiday calendar with national holidays and custom company holidays",
    "leave calculations and attendance tracking correctly exclude these non-working days",
    5, "P0 — Must Have")
sprint_acceptance([
    "The holiday calendar is accessible at /app/hris/holidays.",
    "On org setup, the admin selects a country. National holidays for that country are pre-imported from a static dataset.",
    "Custom company holidays can be added (e.g., 'Founder's Day', 'Company Offsite').",
    "Holidays are year-specific. A 'Copy to Next Year' button duplicates the current year's custom holidays.",
    "Holidays are used by: leave request working-day calculator, attendance auto-close logic, and the Team Calendar.",
])

sprint_story(32, "Team Calendar View",
    "any organisation member",
    "view a calendar showing team members' approved leave, work anniversaries, and company holidays",
    "I can plan my own leave and meetings without conflicts",
    8, "P1 — Should Have")
sprint_acceptance([
    "The Team Calendar at /app/hris/team-calendar shows a monthly calendar view.",
    "Approved leave entries are shown as colored banners spanning the leave date range, color-coded by leave type.",
    "Company and national holidays are shown as a background highlight on the calendar.",
    "Work anniversaries are shown as small icons on the member's anniversary date.",
    "Filter by Department or Team to see only that group's calendar.",
    "Clicking a leave entry shows: employee name, leave type, dates, approval status.",
])
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §9 — PHASE B (SPRINTS 10–23)
# ════════════════════════════════════════════════════════════════════════════
h1("9. Phase B — Feature Depth (Sprints 10–23)")
body("Phase B begins immediately after Phase A is deployed to staging and passes QA. Phase B adds depth and breadth to the platform, expanding the Work OS with advanced views and features, deepening the HRIS with performance reviews and payroll integration, and introducing three major new modules: CRM, ITSM, and advanced reporting. Phase B spans Sprints 10 through 23, representing fourteen two-week sprints or approximately 28 weeks.")

header_table(
    ["Sprint", "Title",                                  "Weeks",        "Team"],
    [
        ["10", "Time Tracking",                          "Weeks 19–20",  "Full-Stack"],
        ["11", "Sprint Management and Backlog",          "Weeks 21–22",  "Full-Stack"],
        ["12", "Gantt Chart and Mind Map Views",         "Weeks 23–24",  "Full-Stack"],
        ["13", "Workload View, Forms, Templates",        "Weeks 25–26",  "Full-Stack"],
        ["14", "Performance Review Cycles",              "Weeks 27–28",  "Full-Stack + HR"],
        ["15", "Onboarding/Offboarding Automation",      "Weeks 29–30",  "Full-Stack"],
        ["16", "Payroll Integration Layer",              "Weeks 31–32",  "Backend + Security"],
        ["17", "CRM Contacts and Companies",             "Weeks 33–34",  "Full-Stack (CRM Team)"],
        ["18", "Pipelines, Deals, Revenue Tracking",    "Weeks 35–36",  "Full-Stack (CRM Team)"],
        ["19", "Lead Scoring and Email Sequences",       "Weeks 37–38",  "Full-Stack (CRM Team)"],
        ["20", "Meeting Scheduling and Activity Log",    "Weeks 39–40",  "Full-Stack (CRM Team)"],
        ["21", "ITSM Service Catalogue and Tickets",     "Weeks 41–42",  "Full-Stack (ITSM Team)"],
        ["22", "Ticket Routing, SLA, Escalation",        "Weeks 43–44",  "Full-Stack (ITSM Team)"],
        ["23", "CSAT, Reporting, Knowledge Base",        "Weeks 45–46",  "Full-Stack (ITSM Team)"],
    ],
    col_widths=[0.6, 3.0, 1.2, 2.3]
)

# ── SPRINT 10 ────────────────────────────────────────────────────────────
sprint_header("B", 10, "Time Tracking", "Weeks 19–20",
"Build the complete Time Tracking feature: a timer widget, manual time entry, time log per task, team-level time reports, and integration with the task time estimate so that actuals vs. estimates can be compared.")

h2("9.1 Sprint 10 — Time Tracking")

sprint_story(33, "Timer Widget",
    "a Work OS user",
    "start and stop a timer against any task directly from the task list or detail panel",
    "time spent on a task is automatically recorded without manual calculation",
    8, "P0 — Must Have")
sprint_acceptance([
    "A play-button (▶) icon appears on each task row in List View and in the Task Detail Panel.",
    "Clicking ▶ starts a timer. The task row shows a running clock (HH:MM:SS). The Slipwise sidebar shows an active timer indicator.",
    "Only one timer can run at a time per user. Starting a new timer pauses the previous one.",
    "Clicking ⏹ stops the timer and creates a TimeEntry record with startTime, endTime, and duration (seconds).",
    "The task's timeTracked field is updated: timeTracked += duration.",
    "Browser tab is closed or user navigates away: timer keeps running (server-side). Client reconnects and resumes the timer display.",
])
sprint_tech([
    "Server-side timer: store {userId, taskId, startTime} in Redis. On stop: compute duration = now - startTime, create TimeEntry, clear Redis key.",
    "Heartbeat: client sends keepalive every 60 seconds. If no heartbeat for 8 hours, auto-stop timer.",
])

sprint_story(34, "Manual Time Entry",
    "a Work OS user",
    "manually log time against a task by specifying the date, hours, and minutes",
    "time that was tracked outside of Slipwise can be recorded accurately",
    5, "P0 — Must Have")
sprint_acceptance([
    "A '+' button in the Task Detail Panel's Time Tracking section opens a 'Log Time' modal.",
    "Modal fields: Date (date picker, defaults to today), Hours, Minutes, Note (optional, max 200 chars).",
    "POST /api/v1/time-entries creates a TimeEntry with type = MANUAL and increments task.timeTracked.",
    "Manual entries are shown in the time log with a 'Manual' badge.",
    "Time entries can be edited or deleted. Editing updates the duration and adjusts task.timeTracked accordingly.",
])

sprint_story(35, "Time Tracking Reports",
    "a manager or HR administrator",
    "view time tracking reports showing how many hours each team member logged per project per week",
    "I can understand where team capacity is being spent and compare actuals to estimates",
    8, "P1 — Should Have")
sprint_acceptance([
    "Time report at /app/work-os/reports/time-tracking shows: Team Member, Week, Hours Logged per Space/List.",
    "Filters: Date Range, Space (multi-select), Team Member (multi-select).",
    "A bar chart visualises daily hours logged for each team member.",
    "Estimate vs. Actual comparison table: for each task with a time estimate, show estimate vs. total time tracked and variance.",
    "Export as CSV with full time entry detail (task, date, duration, note, member).",
])
page_break()

# ── SPRINT 11 ────────────────────────────────────────────────────────────
sprint_header("B", 11, "Sprint Management and Backlog", "Weeks 21–22",
"Build Sprint Management: the ability to create Sprint containers, move tasks into Sprints from the Backlog, start and complete Sprints with rollover handling, and view Sprint velocity reports over time.")

h2("9.2 Sprint 11 — Sprint Management and Backlog")
body("Sprint Management transforms Slipwise Work OS from a general task manager into a tool that engineering teams can use for formal agile development. A Sprint is a time-boxed container within a List or Space. Tasks are pulled from the Backlog into a Sprint, and at the end of the Sprint, incomplete tasks are either rolled over to the next Sprint or returned to the Backlog.")

sprint_story(36, "Create and Manage Sprints",
    "an engineering team lead using Slipwise",
    "create named Sprints with start and end dates within a List or Space and add tasks to them",
    "my team has a clear, time-boxed set of work to commit to each sprint cycle",
    8, "P0 — Must Have")
sprint_acceptance([
    "POST /api/v1/sprints creates a Sprint with: name, listId, startDate, endDate, goal (optional text).",
    "A List can have multiple Sprints (historical and future) but only one ACTIVE Sprint at a time.",
    "POST /api/v1/sprints/:id/start marks the Sprint as ACTIVE. Future sprints are PLANNED. Past are COMPLETED.",
    "Task rows in the List View show a Sprint badge when assigned to a sprint.",
    "Drag-and-drop from Backlog view to a Sprint adds the task to that Sprint.",
    "Sprint planning view shows: Backlog on the left, Sprint tasks on the right, with story point total and team capacity indicator.",
])

sprint_story(37, "Backlog Management",
    "an engineering team lead",
    "view all tasks not yet assigned to a Sprint in a prioritised Backlog list and plan which tasks to include in upcoming Sprints",
    "Sprint planning sessions are efficient and data-driven",
    8, "P0 — Must Have")
sprint_acceptance([
    "The Backlog view is a separate tab within a List: shows all tasks with no Sprint assignment, ordered by priority then creation date.",
    "Tasks in the Backlog can be sorted by Priority, Story Points, or Creation Date.",
    "Bulk-select tasks from the Backlog and assign them to any PLANNED Sprint.",
    "Story point total of selected backlog items is shown in real-time as tasks are checked.",
    "Epic grouping: tasks can be tagged with an Epic label and the Backlog can be grouped by Epic.",
])

sprint_story(38, "Sprint Completion and Rollover",
    "an engineering team lead",
    "complete a Sprint and choose what happens to incomplete tasks",
    "there is a clean audit trail of what was delivered vs. carried over each sprint",
    8, "P0 — Must Have")
sprint_acceptance([
    "POST /api/v1/sprints/:id/complete triggers the Sprint Completion wizard.",
    "The wizard shows: tasks completed (status = CLOSED type), tasks incomplete.",
    "For incomplete tasks, admin selects: Move to Backlog, Move to Next Sprint, or Keep in Completed Sprint (for reference).",
    "After completion, a Sprint Summary report is generated showing: total tasks, completed, incomplete, velocity (story points completed), burn-down chart data.",
    "Sprint velocity is stored and used in the Velocity Trend report for future planning.",
])
page_break()

# ── SPRINT 12 ────────────────────────────────────────────────────────────
sprint_header("B", 12, "Gantt Chart and Mind Map Views", "Weeks 23–24",
"Build two advanced Work OS views: the Gantt Chart (interactive timeline view with dependencies) and the Mind Map view (hierarchical visual structure of tasks).")

h2("9.3 Sprint 12 — Gantt Chart and Mind Map Views")

sprint_story(39, "Gantt Chart View",
    "a project manager",
    "view all tasks in a List on a horizontal timeline (Gantt chart) with dependency arrows",
    "I can see the overall project schedule, critical path, and any scheduling conflicts at a glance",
    13, "P1 — Should Have")
sprint_acceptance([
    "The Gantt View is accessible as a view tab. It shows tasks as horizontal bars on a timeline.",
    "Bar width represents the task's duration (startDate to dueDate).",
    "Dependencies are drawn as arrows between task bars. Finish-to-Start shows an arrow from the end of the predecessor to the start of the successor.",
    "Dragging the left edge of a bar adjusts the start date. Dragging the right edge adjusts the end date. Dragging the bar itself moves both dates.",
    "The timeline header supports zoom levels: Day, Week, Month, Quarter. Keyboard shortcut +/- zooms in/out.",
    "Critical path (the longest chain of dependent tasks) is highlighted in red.",
    "Row grouping by Assignee or by Sprint is available.",
    "Tasks with no start or due date are shown in a side list panel and can be dragged onto the timeline to set dates.",
])
sprint_tech([
    "Use DHTMLX Gantt or build a custom SVG-based Gantt. Custom SVG preferred for design control.",
    "Critical path algorithm: Floyd-Warshall on the task dependency graph. Pre-compute on load.",
])

sprint_story(40, "Mind Map View",
    "a creative or planning user",
    "view tasks in a Mind Map layout where the List is the root node and tasks and subtasks radiate outward as branches",
    "I can plan and review complex projects in a non-linear, visual way",
    8, "P2 — Could Have")
sprint_acceptance([
    "The Mind Map view renders the List name as the root node, with tasks as first-level branches and subtasks as second-level branches.",
    "Nodes are colour-coded by task status.",
    "Clicking a node opens the task detail panel.",
    "New tasks can be created by pressing Enter on a node (adds a sibling) or Tab (adds a child subtask).",
    "The mind map can be panned (click-drag) and zoomed (scroll or pinch).",
    "A layout toggle switches between radial layout and hierarchical (top-down) layout.",
])
page_break()

# ── SPRINT 13 ────────────────────────────────────────────────────────────
sprint_header("B", 13, "Workload View, Forms, and Templates", "Weeks 25–26",
"Build the Workload View (capacity planning), Task Forms (public or internal intake forms that create tasks), and List/Space Templates.")

h2("9.4 Sprint 13 — Workload View, Forms, Templates")

sprint_story(41, "Workload View",
    "a team manager",
    "view a Workload calendar showing how many tasks and hours each team member is assigned to across a time range",
    "I can identify overloaded team members and redistribute work before burnout occurs",
    13, "P1 — Should Have")
sprint_acceptance([
    "The Workload View at /app/work-os/workload shows a matrix: rows are team members, columns are weeks (default: next 4 weeks).",
    "Each cell shows the total time estimate of tasks due that week for that member, as a bar with colour coding: green (< 80% capacity), yellow (80–100%), red (> 100% capacity).",
    "Capacity is configurable per member (default 40 hours/week for full-time, proportional for part-time).",
    "Clicking a cell shows the list of tasks contributing to that member's load that week.",
    "Tasks can be reassigned or rescheduled directly from the workload view by clicking a task and changing the assignee or due date.",
    "Filter by Space, Department, or Team to scope the workload view.",
])

sprint_story(42, "Task Intake Forms",
    "a team administrator",
    "create a Task Form that others can fill in to submit requests which automatically create tasks in a specified List",
    "my team has a structured, consistent way to receive work requests",
    8, "P1 — Should Have")
sprint_acceptance([
    "POST /api/v1/forms creates a Form linked to a List with a name, description, and list of form fields.",
    "Form fields map to task properties: Title (text, required), Description (rich text), Priority (dropdown), Custom Fields (any Custom Field defined on the Space).",
    "Each Form gets a unique public URL (/forms/:formSlug) and can optionally require Slipwise login to submit.",
    "Submitting the form creates a Task in the linked List with the submitted values.",
    "The form submitter receives a confirmation email with a reference number (the task ID).",
    "Form submissions are tracked in a 'Form Responses' tab on the List view showing all submissions.",
])

sprint_story(43, "Space and List Templates",
    "an organisation administrator",
    "create Space or List templates that capture a predefined structure (statuses, custom fields, and optionally starter tasks)",
    "new projects can be bootstrapped in seconds rather than set up from scratch",
    5, "P2 — Could Have")
sprint_acceptance([
    "POST /api/v1/templates creates a template by snapshotting an existing Space or List's configuration.",
    "Template captures: statuses, custom field definitions, views, and optionally a set of template tasks (with relative due dates, e.g., 'Task due 3 days after space creation').",
    "When creating a new Space or List, the user can choose from available templates.",
    "A template library has built-in Slipwise-provided templates: Software Sprint, Marketing Campaign, Client Onboarding, Bug Tracker.",
])
page_break()

# ── SPRINT 14 ────────────────────────────────────────────────────────────
sprint_header("B", 14, "Performance Review Cycles", "Weeks 27–28",
"Build the Performance Review module: review cycle configuration, self-assessment forms, manager review forms, peer review (360-degree) collection, calibration workflow, and review summary reports.")

h2("9.5 Sprint 14 — Performance Review Cycles")
body("Performance reviews are a critical HR process that directly affects employee compensation, promotions, and engagement. The system must support multiple review types (annual, mid-year, probationary, 360), allow HR to configure custom rating scales and competency frameworks, and generate clear summaries that both employees and managers can reference.")

sprint_story(44, "Review Cycle Configuration",
    "an HR administrator",
    "configure a Performance Review Cycle specifying the review type, participants, rating scale, review questions, and timeline",
    "performance reviews are consistently structured across the organisation",
    8, "P0 — Must Have")
sprint_acceptance([
    "POST /api/v1/review-cycles creates a cycle with: name, type (ANNUAL/MIDYEAR/PROBATION/360), reviewees (ALL_ACTIVE / DEPARTMENT / SPECIFIC_LIST), ratingScale (1-3, 1-4, 1-5, or custom labels), selfAssessmentEnabled, managerReviewEnabled, peerReviewEnabled, peerNominationEnabled, startDate, submissionDeadline.",
    "Review questions are configured as a list of competency blocks, each with a title and sub-questions (rating + comment).",
    "Example competency blocks: Communication, Technical Skills, Collaboration, Initiative, Delivery.",
    "On cycle activation, ReviewAssignment records are created for every reviewee-reviewer pair.",
])

sprint_story(45, "Self-Assessment and Manager Review Forms",
    "an employee and a manager",
    "complete a digital self-assessment form and a manager review form respectively for each reviewee",
    "review data is captured consistently and is available for calibration",
    8, "P0 — Must Have")
sprint_acceptance([
    "Each employee sees a 'My Reviews' section listing pending self-assessment forms.",
    "The self-assessment form shows each competency with a rating slider and a text box for comments.",
    "Managers see a 'Team Reviews' section listing all direct reports awaiting manager review.",
    "Manager review form shows the employee's self-ratings as a reference alongside the manager's rating inputs.",
    "Forms support auto-save every 30 seconds.",
    "Submitted forms are locked (read-only) until HR re-opens the cycle for edits.",
    "HR can see submission progress: 'X of Y self-assessments submitted, X of Y manager reviews submitted.'",
])

sprint_story(46, "360-Degree Peer Review",
    "an employee",
    "nominate peers to review me and receive peer review requests from colleagues",
    "my performance review includes diverse perspectives beyond just my manager's view",
    5, "P1 — Should Have")
sprint_acceptance([
    "In cycles with peerNominationEnabled, employees can nominate 3–8 peers to review them.",
    "HR approves the nominations (with a minimum of 3 peer reviews required to be valid).",
    "Nominated peers receive email + in-app notification with a link to their peer review form.",
    "Peer reviews are anonymous (individual ratings are hidden from the reviewee, only averages shown).",
    "A minimum of 3 peer reviews is required before the peer review section is included in the final report.",
])
page_break()

# ── SPRINT 15 ────────────────────────────────────────────────────────────
sprint_header("B", 15, "Onboarding and Offboarding Automation", "Weeks 29–30",
"Build automated onboarding and offboarding workflows that trigger a checklist of tasks, notifications, and automated actions when an employee joins or leaves the organisation.")

h2("9.6 Sprint 15 — Onboarding and Offboarding Automation")
body("Onboarding and offboarding are high-effort, high-risk HR processes. Missed steps in onboarding (forgetting to set up email access, not completing compliance training) create bad first impressions and legal risk. Missed steps in offboarding (not revoking system access) create security vulnerabilities. Slipwise automates both with configurable checklists.")

sprint_story(47, "Onboarding Workflow Builder",
    "an HR administrator",
    "configure an onboarding workflow template with a sequence of tasks, system actions, and notifications triggered when a new employee is added",
    "every new hire goes through the same consistent, complete onboarding experience",
    13, "P0 — Must Have")
sprint_acceptance([
    "POST /api/v1/onboarding-templates creates a template with a list of onboarding steps.",
    "Each step has: title, description, assignee type (NEW_HIRE / MANAGER / HR / IT_ADMIN), dueOffset (days after start date), stepType (TASK / NOTIFICATION / DOCUMENT_UPLOAD / SYSTEM_ACTION).",
    "SYSTEM_ACTION step types: 'Send welcome email', 'Assign to team in Work OS', 'Grant module permissions from template'.",
    "When a new member is added via the Onboarding Wizard, the selected onboarding template is triggered automatically.",
    "All generated tasks appear in the respective assignees' Work OS task lists.",
    "HR has an 'Onboarding Dashboard' showing each active new hire's onboarding progress (% tasks completed).",
])

sprint_story(48, "Offboarding Workflow",
    "an HR administrator",
    "initiate an offboarding workflow when an employee leaves, triggering tasks for equipment return, access revocation, and exit interview",
    "no offboarding steps are missed and the organisation is protected from access risks",
    8, "P0 — Must Have")
sprint_acceptance([
    "Offboarding is initiated from the employee's profile: 'Begin Offboarding' button (requires Admin permission).",
    "A wizard asks: Last Working Day, Reason for Leaving (RESIGNATION/TERMINATION/CONTRACT_END/RETIREMENT), Exit Interview Date.",
    "Offboarding checklist is auto-generated from the configured template. Tasks include: IT equipment return, revoke Slipwise access (scheduled for last working day), exit interview, final payroll calculation, LinkedIn reference offer.",
    "On the last working day: a BullMQ job automatically deactivates the user's account (status = OFFBOARDED, all sessions invalidated).",
    "The offboarded user's data is retained per the organisation's data retention policy (configurable, default 2 years).",
    "All tasks assigned to the departing employee are automatically reassigned to their manager with a notification.",
])
page_break()

# ── SPRINT 16 ────────────────────────────────────────────────────────────
sprint_header("B", 16, "Payroll Integration Layer", "Weeks 31–32",
"Build the Payroll Integration Layer that connects Slipwise HRIS data (salary bands, leave balances, attendance records) to external payroll systems via a standardised export format and direct API integrations with major payroll providers.")

h2("9.7 Sprint 16 — Payroll Integration Layer")
body("Payroll is the most financially sensitive system in an organisation. While Slipwise does not process payroll itself (payroll involves banking, tax compliance, and regulatory approvals that require dedicated licensed platforms), it maintains the HR source data that feeds into payroll systems. This sprint builds the bridge between Slipwise's HRIS and payroll processors.")

sprint_story(49, "Payroll Data Model",
    "the engineering team",
    "have a data model for Compensation (salary bands, effective dates) and Payroll Runs (period, status)",
    "payroll data is stored securely and can be accessed by authorised integrations",
    8, "P0 — Must Have")
sprint_acceptance([
    "CompensationRecord: id, employeeId, currency, grossSalary, salaryBand, effectiveDate, endDate, paymentFrequency (MONTHLY/BIWEEKLY/WEEKLY), createdBy.",
    "All CompensationRecord fields are field-level encrypted at rest (application-layer AES-256 before Prisma write).",
    "Access to compensation records requires the 'payroll:read' permission — separate from general HR read.",
    "PayrollRun: id, orgId, period (YYYY-MM), status (DRAFT/FINALISED/EXPORTED), exportedAt, exportFormat.",
    "PayrollRunItem: id, payrollRunId, employeeId, grossPay, deductions (JSON), netPay, workingDays, paidLeaveDays, unpaidLeaveDays.",
])

sprint_story(50, "Payroll Export",
    "an HR / Finance administrator",
    "generate a payroll export for a given month that includes each employee's gross pay, leave deductions, and attendance summary",
    "I can hand this data to the payroll processor or upload it to our payroll software",
    8, "P0 — Must Have")
sprint_acceptance([
    "POST /api/v1/payroll-runs creates a DRAFT PayrollRun for the selected month.",
    "The system auto-calculates each employee's: Working Days (from attendance records), Unpaid Leave Days (from approved unpaid leave requests), Gross Pay (from latest CompensationRecord), Deductions for unpaid leave (grossPay / workingDaysPerMonth * unpaidLeaveDays).",
    "HR reviews the draft, can make manual adjustments, then clicks Finalise.",
    "Export formats: CSV (generic), QuickBooks IIF, Tally XML (configurable in Settings).",
    "Exported payroll runs are immutable — the export file is archived to S3 with a timestamped filename.",
    "Each payroll export creates an AuditLog record.",
])
page_break()

# ── SPRINT 17 ────────────────────────────────────────────────────────────
sprint_header("B", 17, "CRM — Contacts and Companies", "Weeks 33–34",
"Build the CRM foundation: Contacts and Companies data models, import from CSV, the Contact/Company profile pages with activity timeline, and the preliminary data enrichment hook.")

h2("9.8 Sprint 17 — CRM Contacts and Companies")
body("The CRM module is Slipwise's answer to HubSpot Sales Hub for the SMB market. It must be fast, visual, and tightly integrated with the rest of the platform. When a deal is won, it creates tasks automatically. When a support ticket is raised, it links to the contact's CRM profile. The CRM starts with the foundational entities: Contacts (people) and Companies (organisations the contacts work at).")

sprint_story(51, "Contact Model and CRUD",
    "a sales team member",
    "create, view, edit, and delete Contact records representing individual people in my network",
    "all prospect and customer information is centralised and accessible",
    8, "P0 — Must Have")
sprint_acceptance([
    "A Contact has: firstName, lastName, email[], phone[], linkedInUrl, twitterHandle, jobTitle, company (CompanyId), leadSource, contactStatus (LEAD/PROSPECT/CUSTOMER/CHURNED), tags[], ownerId, createdAt.",
    "POST /api/v1/contacts creates a contact. Email uniqueness enforced within the organisation.",
    "GET /api/v1/contacts returns a paginated list with search, filter (by status, tag, owner, company), and sort.",
    "Contact Profile page shows all fields, plus a rich activity timeline (calls, emails, notes, deal associations, task associations).",
    "Bulk import via CSV: upload a CSV, map columns to Contact fields, preview 5 rows, confirm import. Handles duplicates (match by email) with a merge option.",
])

sprint_story(52, "Company Model and CRUD",
    "a sales team member",
    "create Company records and associate Contacts to them",
    "I can manage relationships at both the individual and organisation level",
    5, "P0 — Must Have")
sprint_acceptance([
    "A Company has: name, domain, industry, size (headcount range), annualRevenue, country, city, website, linkedInUrl, description, tags[], ownerId.",
    "POST /api/v1/companies creates a Company. Domain uniqueness enforced within the org.",
    "Company Profile page shows: company details, associated Contacts list, associated Deals list, total revenue (sum of won deals), activity timeline.",
    "Contacts can be linked to a Company via the company field. One contact can only have one primary company.",
    "Bulk import via CSV supported with same duplicate detection as Contacts.",
])
page_break()

# ── SPRINT 18 ────────────────────────────────────────────────────────────
sprint_header("B", 18, "CRM — Pipelines, Deals, and Revenue Tracking", "Weeks 35–36",
"Build the Deal and Pipeline system: configurable sales pipelines with stages, deal creation and management, deal value and close date tracking, and a revenue dashboard.")

h2("9.9 Sprint 18 — Pipelines, Deals, Revenue Tracking")

sprint_story(53, "Pipeline and Stage Configuration",
    "a sales team lead",
    "create and configure multiple sales pipelines each with custom stages, probability percentages, and deal types",
    "my CRM reflects our actual sales process rather than a generic template",
    5, "P0 — Must Have")
sprint_acceptance([
    "POST /api/v1/pipelines creates a Pipeline with a name and ordered list of Stages.",
    "Each Stage has: name, probability (0–100%), rottenTime (days before deal is considered 'rotten'), color.",
    "Default pipeline 'Sales Pipeline' is created on CRM module activation with stages: Lead, Qualified, Proposal, Negotiation, Won, Lost.",
    "Pipelines can be reordered. Stages within a pipeline can be reordered.",
    "A Stage can be designated as Won (positive final stage) or Lost (negative final stage). Won/Lost stages are terminal — deals in these stages cannot be moved to other stages.",
])

sprint_story(54, "Deal CRUD and Kanban View",
    "a sales team member",
    "create and manage Deals as they progress through the sales pipeline",
    "I have a real-time, visual view of my entire book of business",
    13, "P0 — Must Have")
sprint_acceptance([
    "A Deal has: title, pipelineId, stageId, contactId, companyId, ownerId, value (amount + currency), probability (overridable), expectedCloseDate, actualCloseDate, dealType, lostReason, tags[], createdAt.",
    "The CRM Pipeline View renders deals as cards in a Kanban board — one column per stage. Column header shows stage name and sum of deal values.",
    "Deals can be dragged between stage columns to update their stage.",
    "Deal value total is shown in each column header and a grand total is shown for the pipeline.",
    "Deals with an expectedCloseDate in the past (and not Won/Lost) are highlighted with a 'Past Close Date' warning badge.",
    "Deals that have not had any activity in their rottenTime window are shown with a 'Rotten' 🔴 badge.",
    "Deal filtering: by Owner, Stage, Company, Tag, Expected Close Date range.",
])

sprint_story(55, "Revenue Dashboard",
    "a sales manager",
    "view a revenue dashboard showing pipeline value, forecasted revenue, deal velocity, and won/lost trends",
    "I have the data I need to make informed decisions about the sales team's performance",
    8, "P1 — Should Have")
sprint_acceptance([
    "The Revenue Dashboard at /app/crm/revenue shows: Total Pipeline Value (all open deals), Weighted Forecast (sum of deal value × probability), Deals Closed This Month (won value), Month-Over-Month growth.",
    "A funnel chart shows deal count and value at each pipeline stage.",
    "A line chart shows weekly won deal value over the last 12 months.",
    "Win/Loss analysis table: breakdown by lost reason, deal size, and owner.",
    "All charts are filterable by Date Range, Pipeline, and Team Member.",
])
page_break()

# ── SPRINT 19 ────────────────────────────────────────────────────────────
sprint_header("B", 19, "Lead Scoring and Email Sequences", "Weeks 37–38",
"Build the Lead Scoring engine and the Email Sequence system for CRM outreach automation.")

h2("9.10 Sprint 19 — Lead Scoring and Email Sequences")

sprint_story(56, "Lead Scoring Engine",
    "a sales team lead",
    "configure a lead scoring model that automatically scores contacts based on firmographic and behavioural criteria",
    "my team can prioritise outreach on the highest-potential leads",
    8, "P1 — Should Have")
sprint_acceptance([
    "POST /api/v1/crm/scoring-models creates a scoring model with a list of scoring rules.",
    "Each rule: criterion (field name, e.g. company.size, contactStatus, tag), operator (EQUALS/CONTAINS/GT/LT), value, score (positive or negative integer).",
    "Example rules: company.size = '51-200' → +10 points; leadSource = 'Website' → +5 points; no activity in 30 days → -15 points.",
    "Lead scores are recalculated by a nightly BullMQ job. Score is stored as Contact.leadScore.",
    "In the Contact list view, a score badge (0-100) is shown next to each contact's name.",
    "A 'Hot Leads' filter preset shows contacts with leadScore ≥ 70.",
])

sprint_story(57, "Email Sequence Builder",
    "a BDR or sales team member",
    "create automated email sequences that send a series of templated emails to contacts at scheduled intervals",
    "I can run scalable outbound campaigns without manually sending each email",
    13, "P1 — Should Have")
sprint_acceptance([
    "POST /api/v1/crm/sequences creates a Sequence with: name, goal, steps[].",
    "Each step: delay (days after previous step or enrolment), subject, body (Tiptap rich text with merge tags like {{firstName}}, {{companyName}}), sendTime (09:00 in recipient's timezone by default).",
    "Contacts are enrolled in a Sequence individually or via a bulk action from the Contact list.",
    "Emails are sent via AWS SES at the scheduled time using a BullMQ delayed job.",
    "Auto-reply detection: if a contact replies to any email in the sequence, they are automatically unenrolled.",
    "Sequence analytics: Open Rate, Click Rate, Reply Rate, Unsubscribe Rate per step and per sequence overall.",
    "Unsubscribe link is automatically appended to all sequence emails (CAN-SPAM compliance).",
])
page_break()

# ── SPRINT 20 ────────────────────────────────────────────────────────────
sprint_header("B", 20, "Meeting Scheduling and CRM Activity Log", "Weeks 39–40",
"Build meeting scheduling integrated with CRM contacts, manual and auto-logged activity recording, and the complete CRM activity timeline.")

h2("9.11 Sprint 20 — Meeting Scheduling and Activity Log")

sprint_story(58, "CRM Activity Logging",
    "a sales team member",
    "log calls, emails, meetings, and notes against any Contact or Deal",
    "there is a complete historical record of every customer interaction",
    8, "P0 — Must Have")
sprint_acceptance([
    "Activity types: CALL (duration, outcome, notes), EMAIL (subject, body, direction: inbound/outbound), MEETING (date, duration, attendees, notes, outcome), NOTE (free text), TASK (links to Work OS task).",
    "Activities are logged via a '+' button on any Contact or Deal profile.",
    "All activities appear in the Contact/Deal activity timeline in reverse chronological order.",
    "Activities can be filtered on the timeline by type, date range, and author.",
    "Future: auto-log emails via Google Workspace or Microsoft 365 connector (Sprint 34).",
])

sprint_story(59, "Meeting Scheduler",
    "a sales team member",
    "generate a meeting booking link that prospects can use to book a meeting directly on my calendar",
    "scheduling meetings with prospects is frictionless and requires no back-and-forth emails",
    8, "P1 — Should Have")
sprint_acceptance([
    "Each user can create a Scheduling Link at /app/crm/schedule/:username/:slug.",
    "The link shows available time slots based on the user's configured working hours and calendar availability (Google Calendar integration in Sprint 34; for now, manually configured available slots).",
    "Visitor selects a time slot, enters their name and email, and submits.",
    "A MeetingBooking record is created. Both the host and visitor receive a calendar invite via email with an ICS attachment.",
    "The booking is automatically logged as a MEETING activity on the Contact's profile if the visitor's email matches a Contact in the CRM.",
])
page_break()

# ── SPRINT 21 ────────────────────────────────────────────────────────────
sprint_header("B", 21, "ITSM — Service Catalogue and Ticket Submission", "Weeks 41–42",
"Build the Internal Service Desk: the service catalogue, ticket submission form, ticket list view, and the ticket detail panel. The ITSM module is for internal IT and operational support requests.")

h2("9.12 Sprint 21 — ITSM Service Catalogue and Ticket Submission")
body("The ITSM module transforms Slipwise into the internal service desk for every department. Instead of emailing IT, HR, or Finance for requests, employees open a ticket through the Slipwise service catalogue and the request is automatically routed to the right team. The service catalogue is a directory of all the services each department offers, each service having its own intake form and routing rules.")

sprint_story(60, "Service Catalogue",
    "an IT/HR/Finance department administrator",
    "configure a Service Catalogue listing all internal services my team provides with a name, description, category, and intake form",
    "employees can browse available services and know exactly what to request and how",
    8, "P0 — Must Have")
sprint_acceptance([
    "POST /api/v1/itsm/services creates a Service with: name, description, category (IT/HR/FINANCE/FACILITIES/OTHER), icon, assignedTeamId, defaultPriority, SLA configuration.",
    "The service catalogue homepage at /app/itsm shows all available services grouped by category.",
    "Each service card shows: name, description, icon, estimated resolution time (from SLA config).",
    "Services can be hidden (draft mode) until ready for employee access.",
    "Seed catalogue includes: IT Hardware Request, Software License Request, VPN Access, Password Reset, New Hire IT Setup, Leave Policy Query, Expense Reimbursement.",
])

sprint_story(61, "Ticket Submission",
    "any employee",
    "submit a support ticket for any service from the service catalogue by filling in a context-aware intake form",
    "my request is immediately routed to the right team with all necessary information",
    8, "P0 — Must Have")
sprint_acceptance([
    "Clicking a service in the catalogue opens a Ticket Submission modal with the service's intake form.",
    "Intake form fields are configured per service (title, description, attachments, priority, and custom fields).",
    "POST /api/v1/itsm/tickets creates a Ticket in OPEN status and triggers routing logic.",
    "Submitter receives an immediate email acknowledgement with a ticket number (e.g., ITS-00234).",
    "Ticket list at /app/itsm/my-tickets shows all tickets the user has submitted with status badges.",
    "Tickets support a public comment thread for back-and-forth between the submitter and agent.",
])
page_break()

# ── SPRINT 22 ────────────────────────────────────────────────────────────
sprint_header("B", 22, "ITSM — Ticket Routing, SLA, and Escalation", "Weeks 43–44",
"Build intelligent ticket routing (manual and rule-based), SLA timer enforcement with breach alerts, and an escalation policy engine.")

h2("9.13 Sprint 22 — Ticket Routing, SLA, Escalation")

sprint_story(62, "Ticket Routing Rules",
    "an ITSM administrator",
    "configure routing rules that automatically assign incoming tickets to the correct agent or team based on the service, priority, keywords, or submitter's department",
    "tickets reach the right person without manual sorting",
    8, "P0 — Must Have")
sprint_acceptance([
    "Routing rules are configured per Service or globally at /app/itsm/settings/routing.",
    "Rule conditions: service equals X, priority equals HIGH, submitter department equals SALES, title contains 'VPN'.",
    "Rule actions: assign to specific agent, assign to team (round-robin or least-loaded), set priority.",
    "Rules are evaluated in order (priority order). First matching rule wins.",
    "If no rule matches, ticket is placed in an unassigned queue visible to ITSM admins.",
    "Round-robin assignment: track a last_assigned cursor per team to distribute evenly.",
])

sprint_story(63, "SLA Enforcement",
    "an ITSM team lead",
    "configure SLA policies per ticket priority defining first response and resolution time targets, with alerts when approaching or breaching SLA",
    "the team's service commitments to employees are tracked and enforced",
    8, "P0 — Must Have")
sprint_acceptance([
    "SLA Policy has: priority (P1/P2/P3/P4), firstResponseTarget (minutes), resolutionTarget (minutes), businessHoursOnly (boolean).",
    "When a ticket is created, the SLA clock starts. If businessHoursOnly, the clock only ticks during configured business hours.",
    "SLA timers are tracked by a BullMQ delayed job that checks SLA status every 5 minutes.",
    "At 75% of the SLA time elapsed: the assigned agent receives a 'SLA Warning' notification.",
    "At 100% (breach): the ticket gets a red 'SLA Breached' badge. The agent and their team lead receive a breach alert.",
    "The SLA status (IN_TIME/BREACHED) is recorded on the ticket for reporting.",
])

sprint_story(64, "Escalation Policy Engine",
    "an ITSM administrator",
    "configure escalation policies that automatically reassign or notify senior agents when a ticket meets certain conditions",
    "critical issues are never stuck waiting with an unavailable agent",
    8, "P1 — Should Have")
sprint_acceptance([
    "Escalation triggers: SLA breach, ticket open for > N hours, priority = P1 unacknowledged for > 30 minutes.",
    "Escalation actions: reassign to escalation agent, add watcher, send email to team lead, change priority to P1.",
    "Multiple escalation tiers: Tier 1 (15 mins) → Tier 2 (30 mins) → Tier 3 Manager (1 hour).",
    "Escalation events are recorded in the ticket activity log.",
    "A 'Manually Escalate' button is available to agents for immediate escalation outside the policy schedule.",
])
page_break()

# ── SPRINT 23 ────────────────────────────────────────────────────────────
sprint_header("B", 23, "ITSM — CSAT, Reporting, and Knowledge Base", "Weeks 45–46",
"Build CSAT collection after ticket resolution, ITSM reporting dashboards, and a searchable Knowledge Base where agents document solutions to common issues.")

h2("9.14 Sprint 23 — CSAT, Reporting, Knowledge Base")

sprint_story(65, "CSAT Collection",
    "an ITSM administrator",
    "automatically send a CSAT survey to ticket submitters after their ticket is resolved",
    "the team can measure and improve service quality over time",
    5, "P0 — Must Have")
sprint_acceptance([
    "On ticket RESOLVED: 24 hours later, a CSAT email is sent to the submitter via AWS SES.",
    "CSAT email contains: a 1–5 star rating (click-to-rate, no login required) and an optional comment box.",
    "The CSAT response is linked to the ticket and displayed on the ticket record.",
    "Agents can see their own average CSAT score on their profile.",
    "CSAT surveys not responded to within 7 days expire and are counted as 'No Response'.",
])

sprint_story(66, "Knowledge Base",
    "an ITSM agent",
    "create, categorise, and publish Knowledge Base articles documenting solutions to common issues",
    "submitters can self-serve common issues and agents can close repetitive tickets faster",
    8, "P1 — Should Have")
sprint_acceptance([
    "Knowledge Base at /app/itsm/kb supports: Articles (rich text, Tiptap), Categories (tree), Tags.",
    "Articles have a DRAFT / PUBLISHED / ARCHIVED status workflow.",
    "The ticket submission modal shows a 'Related Articles' section suggesting articles matching the ticket title keywords.",
    "If a submitter clicks an article and it resolves their issue, they can click 'This solved my problem' — which closes their ticket without submitting.",
    "Article view count, helpful vote count, and not-helpful vote count are tracked.",
    "ITSM agents can 'Attach' a KB article to a resolved ticket, adding it to the Related Articles index for that service.",
])
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §10 — PHASE C (SPRINTS 24–38)
# ════════════════════════════════════════════════════════════════════════════
h1("10. Phase C — Intelligence and Scale (Sprints 24–38)")
body("Phase C introduces the most advanced capabilities of Slipwise: Customer Experience lifecycle management, OKR and Goals tracking, a fully configurable Analytics Studio, a real-time collaborative Docs and Wiki system, a no-code Visual Automation Builder, an AI writing and intelligence layer, external integrations (Google Workspace, Slack, Zoom), a Mobile Progressive Web App, and a final hardening and public launch sprint. Phase C spans Sprints 24 through 38.")

header_table(
    ["Sprint", "Title",                                     "Weeks",        "Team"],
    [
        ["24", "CX — Health Scoring and Onboarding Plays",  "Weeks 47–48",  "Full-Stack (CX)"],
        ["25", "Goals and OKR Engine",                      "Weeks 49–50",  "Full-Stack"],
        ["26", "Analytics Studio",                          "Weeks 51–52",  "Full-Stack + Data"],
        ["27", "Docs and Wiki — Foundation",                "Weeks 53–54",  "Full-Stack"],
        ["28", "Docs — Embeds, Collab, Templates",          "Weeks 55–56",  "Full-Stack"],
        ["29", "Visual Automation Builder",                 "Weeks 57–58",  "Full-Stack + Backend"],
        ["30", "Trigger Library and Integration Hooks",     "Weeks 59–60",  "Backend"],
        ["31", "Zapier / Webhook External Automation",      "Weeks 61–62",  "Backend"],
        ["32", "AI Writing and Summarisation",              "Weeks 63–64",  "Full-Stack + AI"],
        ["33", "AI Predictive Analytics and Insights",      "Weeks 65–66",  "Backend + AI"],
        ["34", "Google Workspace Integration",              "Weeks 67–68",  "Backend"],
        ["35", "Slack, Zoom, and Calendar Sync",            "Weeks 69–70",  "Backend"],
        ["36", "Mobile PWA — Foundation",                   "Weeks 71–72",  "Full-Stack (Mobile)"],
        ["37", "Mobile — Offline and Push Notifications",   "Weeks 73–74",  "Full-Stack (Mobile)"],
        ["38", "Hardening, Audit, and Public Launch",       "Weeks 75–76",  "All Teams"],
    ],
    col_widths=[0.6, 3.0, 1.2, 2.3]
)

# ── SPRINT 24 ────────────────────────────────────────────────────────────
sprint_header("C", 24, "CX — Health Scoring and Onboarding Playbooks", "Weeks 47–48",
"Build the Customer Experience module: customer account health scoring from multiple data signals, automated health score alerts, and Onboarding Playbooks that guide CS teams through structured customer onboarding milestones.")

h2("10.1 Sprint 24 — CX Health Scoring and Onboarding Playbooks")
body("Customer Experience (CX) or Customer Success is the function responsible for retaining and expanding revenue from existing customers. The core tool of a CS team is the health score — a composite metric that tells the CS manager which accounts are at risk before they churn. Slipwise CX is built to replace tools like Gainsight and Vitally, which are prohibitively expensive for SMB companies.")

sprint_story(67, "Customer Account and Health Score Model",
    "a Customer Success manager",
    "have a CX Account record for each customer linked to their CRM Company, with a composite health score automatically calculated",
    "I can see which customers are healthy, at risk, or churning at a glance",
    13, "P0 — Must Have")
sprint_acceptance([
    "CXAccount: id, companyId (linked to CRM Company), csmOwnerId, accountStage (ONBOARDING/ADOPTED/RENEWAL/CHURNED), mrr, contractRenewalDate, healthScore (0–100), healthTrend (UP/STABLE/DOWN), lastHealthCalcAt.",
    "Health score formula is configurable per organisation. Default formula weights: Product Usage (30%), Support Ticket Volume (20%), NPS Response (20%), CRM Activity Recency (15%), Contract Renewal Proximity (15%).",
    "Health score is recalculated nightly by a BullMQ job. Score change of ±10 in a week triggers an alert.",
    "The CX Dashboard at /app/cx shows all customer accounts in a list with health score badges, MRR, renewal date, and CSM owner.",
    "Colour coding: Green (70–100), Yellow (40–69), Red (0–39).",
    "A 'Red Accounts' filter shows all accounts scoring below 40 — the daily starting point for CS work.",
])
sprint_tech([
    "Health score components are fetched from different modules (support ticket count from ITSM, NPS from survey table, activity from CRM). A background aggregation job computes these and stores the component scores in a CXHealthScoreBreakdown table for transparency.",
])

sprint_story(68, "Onboarding Playbook",
    "a Customer Success manager",
    "assign an Onboarding Playbook to a new customer account and track milestone completion",
    "every new customer goes through the same structured, proven onboarding experience",
    8, "P0 — Must Have")
sprint_acceptance([
    "POST /api/v1/cx/playbooks creates a Playbook template with milestones. Each milestone: name, description, tasks[], dueOffsetDays (from account creation), ownerRole (CSM/CUSTOMER/PRODUCT).",
    "When a CXAccount is created, the CS manager selects a Playbook. PlaybookAssignment records are created.",
    "Milestones and tasks appear in the CS manager's Work OS task list automatically.",
    "Customer-facing milestones generate a shared checklist accessible to the customer via a secure URL.",
    "The CXAccount profile shows Playbook progress as a visual timeline with milestone status badges.",
    "A 'Playbook Completion' report shows average days to complete each milestone across all accounts for identifying bottlenecks.",
])
page_break()

# ── SPRINT 25 ────────────────────────────────────────────────────────────
sprint_header("C", 25, "Goals and OKR Engine", "Weeks 49–50",
"Build a full OKR (Objectives and Key Results) system with company-level, team-level, and individual objectives, measurable key results with progress tracking, and an OKR review cycle workflow.")

h2("10.2 Sprint 25 — Goals and OKR Engine")

sprint_story(69, "Objective and Key Result CRUD",
    "any manager or senior individual contributor",
    "create Objectives with supporting Key Results, set timeframes, and link them to teams or the organisation",
    "every goal is visible, measurable, and tied to the broader company direction",
    8, "P0 — Must Have")
sprint_acceptance([
    "Objective: id, title, description, owner (userId or teamId), parentObjectiveId (nullable, for cascading OKRs), timeframe (Q1/Q2/Q3/Q4 + year), status (ON_TRACK/AT_RISK/BEHIND/COMPLETED).",
    "Key Result: id, objectiveId, title, type (PERCENTAGE/NUMERIC/MILESTONE/CURRENCY), startValue, targetValue, currentValue, unit (%, $, '#', etc.), dueDate.",
    "Objective progress = average of Key Result progress percentages.",
    "Key Results are updated by the owner manually with a check-in (currentValue + update comment).",
    "OKR alignment tree: Objectives can have parent Objectives, creating a cascade from Company → Department → Team → Individual.",
    "Check-in history is tracked per Key Result showing all value updates with timestamps and comments.",
])

sprint_story(70, "OKR Dashboard",
    "any organisation member",
    "view the OKR dashboard showing all active objectives at every level with progress indicators",
    "everyone can see how the company is progressing toward its goals",
    8, "P0 — Must Have")
sprint_acceptance([
    "OKR Dashboard at /app/goals shows: Company Objectives → Department Objectives → Team Objectives in a tree view.",
    "Each objective shows a circular progress indicator, status badge, owner avatar, and timeframe label.",
    "A 'My OKRs' tab shows only objectives owned by or contributed to by the current user.",
    "Weekly Check-in reminder notifications are sent every Monday morning to all OKR owners with outstanding updates.",
    "End-of-quarter OKR Review: a review window is opened where owners mark Key Results as achieved/missed with a retrospective note.",
])
page_break()

# ── SPRINT 26 ────────────────────────────────────────────────────────────
sprint_header("C", 26, "Analytics Studio", "Weeks 51–52",
"Build the Analytics Studio: a configurable, drag-and-drop dashboard builder that reads live data from all Slipwise modules and allows users to create custom charts, tables, and KPI cards.")

h2("10.3 Sprint 26 — Analytics Studio")

sprint_story(71, "Dashboard Builder",
    "a business analyst or manager",
    "create a custom dashboard by selecting data sources from any Slipwise module and adding chart, table, or KPI card widgets",
    "I have a single pane of glass for business intelligence without needing external tools",
    13, "P1 — Should Have")
sprint_acceptance([
    "Dashboard builder at /app/analytics/dashboards/new uses a drag-and-drop canvas (react-grid-layout).",
    "Widget types: KPI Card (single metric with trend arrow), Bar Chart, Line Chart, Pie Chart, Data Table, Funnel Chart.",
    "Each widget has a data source selector. Data sources: Work OS (task counts, completion rates, velocity), HRIS (headcount, leave utilisation, attrition), CRM (pipeline value, win rate, deal velocity), ITSM (ticket volume, resolution time, SLA breach rate), CX (average health score, churn rate, NRR).",
    "Each widget has configurable filters (date range, department, team, owner).",
    "Dashboards can be shared with specific team members or made visible to the entire organisation.",
    "Scheduled email delivery: dashboards can be emailed as PDF/PNG snapshots on a schedule (daily/weekly/monthly).",
])

sprint_story(72, "Pre-Built Report Library",
    "a manager without analytics expertise",
    "access a library of pre-built reports for common use cases without building from scratch",
    "I get immediate value from the analytics module on day one",
    5, "P1 — Should Have")
sprint_acceptance([
    "Pre-built reports include: Team Productivity Report (tasks completed per member per week), Leave Utilisation Report, CRM Pipeline Health Report, ITSM Agent Performance Report, CX Churn Risk Report.",
    "Each pre-built report is filterable by date range and department.",
    "Any pre-built report can be cloned and customised.",
    "Reports can be exported as: CSV (raw data), PDF (formatted report), or PNG (chart image).",
])
page_break()

# ── SPRINT 27 ────────────────────────────────────────────────────────────
sprint_header("C", 27, "Docs and Wiki — Foundation", "Weeks 53–54",
"Build the Docs and Wiki module foundation: the block-based document editor (Tiptap), document hierarchy (nested pages), version history, and document permissions.")

h2("10.4 Sprint 27 — Docs and Wiki Foundation")

sprint_story(73, "Block-Based Document Editor",
    "any Slipwise member with Docs access",
    "create and edit rich documents using a block-based editor that supports headings, paragraphs, lists, code blocks, images, and tables",
    "I can create professional documentation without needing external tools like Notion or Confluence",
    13, "P0 — Must Have")
sprint_acceptance([
    "Document editor uses Tiptap v2 with the following block types: Paragraph, Heading (H1/H2/H3), Bulleted List, Numbered List, Task List (checkboxes), Code Block (with syntax highlighting via lowlight), Image (upload or URL), Table (resizable, sortable columns), Blockquote, Divider, Callout (info/warning/danger box).",
    "Slash command menu: typing '/' in the editor shows a searchable list of all block types.",
    "Document content is stored as ProseMirror JSON in the Document.content column.",
    "Documents are saved automatically every 5 seconds after any change.",
    "Full-text search over document content using PostgreSQL tsvector.",
])

sprint_story(74, "Document Hierarchy and Navigation",
    "a team or organisation wiki editor",
    "organise documents into a nested tree of pages with a sidebar navigation panel",
    "the wiki feels as well-organised and navigable as a real documentation site",
    8, "P0 — Must Have")
sprint_acceptance([
    "Documents have: id, title, parentDocumentId (nullable), spaceId (optional — docs can be global or space-scoped), createdBy, lastEditedBy, isPublic, archivedAt.",
    "The Docs sidebar shows a collapsible tree of all documents the user can access.",
    "Documents can be drag-dropped to re-parent them in the tree.",
    "Breadcrumb navigation shows the current document's position in the hierarchy.",
    "Clicking the '+' icon next to any document creates a sub-page underneath it.",
    "Archived documents are hidden from the navigation but accessible via a separate 'Archive' section.",
])

sprint_story(75, "Version History",
    "a document editor",
    "view the version history of a document and restore any previous version",
    "I can recover from accidental deletions and see who changed what and when",
    5, "P1 — Should Have")
sprint_acceptance([
    "Document versions are created every time a document is saved with changes (max 1 version per 5-minute window per user to avoid excessive storage).",
    "Version history panel shows: timestamp, author, word count change, and a diff preview (added/removed text highlighted).",
    "Clicking 'Restore this version' creates a new version with the old content — it does not overwrite history.",
    "Version retention: last 100 versions or 90 days, whichever is smaller. Older versions are pruned by a nightly cleanup job.",
])
page_break()

# ── SPRINT 28 ────────────────────────────────────────────────────────────
sprint_header("C", 28, "Docs — Embeds, Real-Time Collaboration, Templates", "Weeks 55–56",
"Add embedded content (YouTube, Figma, Google Docs, Loom, etc.), real-time multi-user collaborative editing with cursor presence, and a document template library.")

h2("10.5 Sprint 28 — Docs Embeds, Collaboration, Templates")

sprint_story(76, "Real-Time Collaborative Editing",
    "multiple team members working on the same document simultaneously",
    "see each other's cursors and edits in real-time without conflicts",
    "we can co-author documentation as a team without locking or version conflicts",
    13, "P1 — Should Have")
sprint_acceptance([
    "Collaborative editing is powered by Yjs CRDT library with a Hocuspocus WebSocket server for conflict-free real-time sync.",
    "Each collaborator's cursor is shown as a coloured caret with their name.",
    "Active collaborators are shown as avatar bubbles in the document header.",
    "Offline editing: changes are queued locally and synced when the connection is restored.",
    "All changes are merged without conflicts using Yjs's CRDT algorithms.",
    "The Hocuspocus server persists the Yjs document state to PostgreSQL (Y.js binary blob) as the source of truth.",
])

sprint_story(77, "Rich Embeds",
    "a document editor",
    "embed external content (YouTube videos, Figma designs, Loom recordings, Google Sheets) inside a document using a URL",
    "documentation can include live, interactive external content rather than just static screenshots",
    5, "P1 — Should Have")
sprint_acceptance([
    "The '/embed' slash command opens a URL input. Pasting a URL from a supported provider renders an embedded iframe.",
    "Supported embed providers: YouTube, Vimeo, Loom, Figma, Google Docs, Google Sheets, Airtable, CodePen, Miro.",
    "Embeds use the oEmbed standard where available; fallback to iframe for others.",
    "Embed height is resizable by dragging the bottom edge of the embed block.",
    "A 'Link preview' block type shows a title/description/image card for any URL (fetched from the URL's OpenGraph tags).",
])
page_break()

# ── SPRINT 29 ────────────────────────────────────────────────────────────
sprint_header("C", 29, "Visual Automation Builder", "Weeks 57–58",
"Build the no-code Visual Automation Builder: a flow-chart interface where users can create trigger-action automations that run across any Slipwise module.")

h2("10.6 Sprint 29 — Visual Automation Builder")

sprint_story(78, "Automation Flow Canvas",
    "a power user or administrator",
    "create automations using a visual drag-and-drop flow builder with trigger nodes, condition nodes, and action nodes",
    "I can automate repetitive cross-module workflows without writing code",
    13, "P1 — Should Have")
sprint_acceptance([
    "Automation builder at /app/automation/new uses a node-based canvas (React Flow library).",
    "Node types: TRIGGER (the event that starts the automation), CONDITION (IF branch), DELAY (wait N hours/days), ACTION (what to do), SPLIT (parallel branches).",
    "Connections between nodes are drawn as arrows indicating flow direction.",
    "Each automation can have exactly one Trigger node and multiple downstream nodes.",
    "Automations can be enabled/disabled with a toggle. Disabled automations do not fire.",
    "An execution log shows the last 100 runs: trigger event, timestamp, result (SUCCESS/ERROR/SKIPPED), and per-node execution detail.",
])

sprint_story(79, "Trigger and Action Library (Internal)",
    "a power user building an automation",
    "select from a library of triggers and actions covering Work OS, HRIS, CRM, and ITSM",
    "I can automate meaningful cross-module workflows out of the box",
    13, "P1 — Should Have")
sprint_acceptance([
    "Triggers: Task Status Changed, Task Created, Deal Stage Changed, Deal Won, Ticket Created, Ticket SLA Breached, Leave Request Submitted, Employee Added, Form Submitted, Date/Time Trigger (scheduled).",
    "Actions: Create Task, Update Task, Send Notification (in-app), Send Email, Create Deal, Update Deal Stage, Create Ticket, Send Webhook (to external URL), Add to CRM Sequence, Update Custom Field.",
    "Conditions: field value equals/contains/greater than, user department equals, tag contains.",
    "Example automation: 'When a Deal is WON → Create a Work OS task list from the Client Onboarding template → Send Slack notification to CS team → Create CX Account record.'",
    "All automations run in the background via BullMQ workers. No user waits for automation execution.",
])
page_break()

# ── SPRINT 30 ────────────────────────────────────────────────────────────
sprint_header("C", 30, "Trigger Library and Integration Hooks", "Weeks 59–60",
"Expand the automation trigger library with advanced cross-module triggers, build the integration hook system, and introduce scheduled (cron) automation triggers.")

h2("10.7 Sprint 30 — Trigger Library and Integration Hooks")

sprint_story(80, "Scheduled Triggers",
    "an administrator",
    "create automations that run on a schedule — daily, weekly, monthly, or at a specific date and time",
    "I can automate recurring processes like weekly report generation or monthly leave accrual reminders",
    5, "P1 — Should Have")
sprint_acceptance([
    "Scheduled Trigger configuration: frequency (ONCE/DAILY/WEEKLY/MONTHLY), time, day of week (for weekly), day of month (for monthly).",
    "Scheduled automations are managed by BullMQ cron jobs. Cron schedule is computed from the trigger configuration.",
    "Each scheduled run creates an AutomationRun log record.",
    "Scheduled automations can be paused without deleting them.",
])

sprint_story(81, "Incoming Webhooks",
    "a developer or administrator",
    "create an Incoming Webhook URL that external systems can POST data to, triggering a Slipwise automation",
    "external tools that are not natively integrated can still trigger Slipwise workflows",
    8, "P2 — Could Have")
sprint_acceptance([
    "POST /api/v1/automation/webhooks creates an IncomingWebhook with a unique URL slug and a secret token for request signature validation.",
    "External systems POST JSON to the webhook URL with the signature in the X-Slipwise-Signature header.",
    "The webhook trigger node in the automation builder lets users map webhook payload fields to automation variables.",
    "Webhook delivery logs show the last 100 incoming requests with payload and processing result.",
])
page_break()

# ── SPRINT 31 ────────────────────────────────────────────────────────────
sprint_header("C", 31, "Zapier and Outgoing Webhook Automation", "Weeks 61–62",
"Build Zapier integration (native Slipwise app in the Zapier app directory) and robust outgoing webhook infrastructure for connecting Slipwise to hundreds of external tools.")

h2("10.8 Sprint 31 — Zapier and Webhook External Automation")

sprint_story(82, "Outgoing Webhooks",
    "an administrator",
    "configure outgoing webhooks that send event data from Slipwise to external URLs when specific events occur",
    "external tools and custom systems can react to Slipwise events without polling the API",
    8, "P1 — Should Have")
sprint_acceptance([
    "POST /api/v1/webhooks creates an OutgoingWebhook with: url, events[] (e.g., 'task.created', 'deal.won'), secret, isActive.",
    "When a subscribed event occurs, a BullMQ job sends an HTTP POST to the webhook URL within 10 seconds.",
    "Request body: { event, timestamp, organisationId, data: {...eventData} }.",
    "HMAC-SHA256 signature in X-Slipwise-Signature header for receiver verification.",
    "Retry policy: exponential backoff up to 3 retries on non-2xx responses.",
    "Delivery log shows: timestamp, event, response status, response body, latency.",
])
page_break()

# ── SPRINT 32 ────────────────────────────────────────────────────────────
sprint_header("C", 32, "AI Writing and Summarisation Layer", "Weeks 63–64",
"Integrate Google Gemini to power AI writing assistance, document summarisation, task generation from meeting notes, and email composition in the CRM.")

h2("10.9 Sprint 32 — AI Writing and Summarisation")

sprint_story(83, "AI Writing Assistant in Docs",
    "a document editor",
    "use an AI writing assistant to draft sections, improve clarity, summarise long content, or translate text",
    "my documentation quality improves and writing takes less time",
    13, "P1 — Should Have")
sprint_acceptance([
    "In the document editor, pressing Space on an empty block shows an AI icon. Clicking it opens the AI command palette.",
    "AI commands: 'Draft from prompt' (generate a section from a description), 'Improve writing' (rewrite selection for clarity), 'Summarise' (condense selection into a summary), 'Expand' (elaborate on a brief bullet), 'Fix grammar', 'Translate to…' (supports 10 major languages), 'Make shorter / Make longer'.",
    "For 'Draft from prompt': user types a description (e.g., 'Write a quarterly update for the engineering team') and AI generates a multi-paragraph draft.",
    "AI-generated content is shown with a light purple background and an 'AI Generated' label. User can Accept, Discard, or Edit.",
    "All AI prompts and responses are logged (sans PII) for quality monitoring. Company data is never used for model training (enforced by API configuration).",
])
sprint_tech([
    "API: POST /api/v1/ai/write. Server-side call to Gemini 1.5 Pro API with the document context as system prompt.",
    "Use streaming response: Next.js route uses ReadableStream to stream the AI response to the client token by token.",
])

sprint_story(84, "AI Task Generation from Notes",
    "a project manager or team lead",
    "paste meeting notes or a requirements description and have AI extract and create tasks from it",
    "converting meeting outcomes to actionable tasks takes seconds instead of minutes",
    5, "P1 — Should Have")
sprint_acceptance([
    "In any List, a 'Generate Tasks with AI' button opens a text input modal.",
    "User pastes meeting notes or a brief project description.",
    "AI returns a structured list of tasks with: title, priority suggestion, and brief description.",
    "User reviews the generated tasks (can edit, remove, or add tasks) and clicks 'Create All' to bulk-create them.",
    "Tasks are created in the current List with status = first status and reporter = current user.",
])
page_break()

# ── SPRINT 33 ────────────────────────────────────────────────────────────
sprint_header("C", 33, "AI Predictive Analytics and Insights", "Weeks 65–66",
"Extend the AI layer with predictive features: deal win probability prediction, CX churn risk prediction, and the AI Insights feed.")

h2("10.10 Sprint 33 — AI Predictive Analytics and Insights")

sprint_story(85, "Deal Win Probability Prediction",
    "a sales manager",
    "see an AI-calculated probability of winning each open deal based on deal characteristics and historical patterns",
    "I can focus coaching effort on the highest-impact deals",
    8, "P2 — Could Have")
sprint_acceptance([
    "For each open deal, an AI_WIN_PROBABILITY field is calculated by calling a Gemini function-calling prompt that analyses: deal age, value, number of activities, last activity recency, stage, and the organisation's historical win/loss data.",
    "The AI probability is shown as a secondary badge on each deal card in the Pipeline View.",
    "If the AI probability differs from the manual probability by more than 20 points, a flag icon is shown prompting review.",
    "Predictions are recalculated nightly and whenever a significant deal activity is logged.",
])

sprint_story(86, "AI Insights Feed",
    "any manager or executive",
    "see a curated AI Insights feed on my home dashboard with proactive alerts and anomaly detections",
    "I surface issues and opportunities I would otherwise miss in a sea of data",
    8, "P2 — Could Have")
sprint_acceptance([
    "The Home Dashboard includes an 'AI Insights' sidebar card showing up to 5 proactive insights.",
    "Insight types: 'Deal XYZ has been in the Proposal stage for 21 days without activity — consider follow-up', 'Customer ABC's health score dropped 15 points this week', 'Sprint velocity is 20% below the team's 4-sprint average', 'Employee John has 12 days of leave expiring unused this quarter'.",
    "Each insight has an action button (e.g., 'Log Activity', 'View Account', 'View Sprint Report').",
    "Insights are generated nightly by a Gemini analysis job that queries the operational database for anomalies.",
])
page_break()

# ── SPRINT 34 ────────────────────────────────────────────────────────────
sprint_header("C", 34, "Google Workspace Integration", "Weeks 67–68",
"Build native Google Workspace integration: Gmail email sync for CRM activity logging, Google Calendar two-way sync for meetings and leave, and Google Drive attachment picker.")

h2("10.11 Sprint 34 — Google Workspace Integration")

sprint_story(87, "Google Calendar Two-Way Sync",
    "any Slipwise user",
    "connect my Google Calendar account to Slipwise so that meetings booked via Slipwise CRM appear in Google Calendar and vice versa",
    "I have one calendar as my source of truth",
    8, "P1 — Should Have")
sprint_acceptance([
    "OAuth 2.0 flow to connect Google account at /app/settings/integrations/google.",
    "Meetings created in Slipwise CRM Scheduler: automatically created in the user's Google Calendar with attendees and ICS data.",
    "Google Calendar events: optionally sync to Slipwise (user configures which calendars to sync). Synced events appear as blocked time in the Workload View.",
    "Leave approved in Slipwise: automatically adds an 'Out of Office' event to the employee's Google Calendar for the leave dates.",
    "Sync uses Google Calendar Push Notifications (webhooks) for near-real-time updates.",
])

sprint_story(88, "Gmail CRM Activity Auto-Log",
    "a sales team member",
    "have emails sent to or received from a CRM contact automatically logged as activities on that contact's timeline",
    "I never miss an interaction in my CRM history because I forgot to manually log it",
    8, "P1 — Should Have")
sprint_acceptance([
    "After connecting Google account, user enables Gmail sync in settings.",
    "A Gmail push notification (Pub/Sub) triggers processing of new emails where the sender or recipient matches a CRM Contact's email.",
    "Matching emails are logged as EMAIL activity on the Contact and linked Deal.",
    "Email body is truncated to 5000 characters. Subject and sender/recipient are always stored.",
    "Users can 'Exclude' specific email threads from CRM logging via a Gmail label.",
    "Privacy: only emails where the Slipwise user is the sender or recipient are synced — not all Gmail inbox.",
])
page_break()

# ── SPRINT 35 ────────────────────────────────────────────────────────────
sprint_header("C", 35, "Slack, Zoom, and Calendar Sync", "Weeks 69–70",
"Build Slack integration (notifications and slash commands), Zoom meeting creation from CRM, and Microsoft Calendar sync as an alternative to Google.")

h2("10.12 Sprint 35 — Slack, Zoom, Calendar Sync")

sprint_story(89, "Slack Integration — Notifications",
    "any Slipwise user who uses Slack",
    "receive Slipwise notifications in a designated Slack channel or as direct messages",
    "I don't need to keep Slipwise open to stay on top of important updates",
    8, "P1 — Should Have")
sprint_acceptance([
    "OAuth 2.0 connection at /app/settings/integrations/slack.",
    "Notification types routable to Slack: Task assigned to me, Task comment mentioning me, Leave request submitted (for managers), ITSM ticket assigned, Deal won, CX health score alert.",
    "Each notification type can be enabled/disabled independently.",
    "Workspace-level Slack: HR announcements and leave approval notifications can be posted to a configured #hr-updates channel.",
    "Slack messages are formatted using Slack Block Kit with action buttons (e.g., 'Approve Leave' button that calls the Slipwise API directly from Slack).",
])

sprint_story(90, "Zoom Meeting Creation",
    "a sales team member",
    "create a Zoom meeting directly from a Slipwise CRM deal or meeting booking and have the Zoom link auto-inserted",
    "setting up video calls for prospects is frictionless and always has the right link",
    5, "P2 — Could Have")
sprint_acceptance([
    "OAuth 2.0 connection at /app/settings/integrations/zoom.",
    "When creating a MEETING activity or a scheduled meeting booking, a 'Create Zoom Meeting' toggle is available.",
    "On toggle: Slipwise calls Zoom API to create a meeting and stores the join URL and meeting ID.",
    "The Zoom join URL is included in the calendar invite email sent to attendees.",
    "After the meeting ends, the Zoom recording URL (if cloud recording enabled) is automatically attached to the meeting activity log.",
])
page_break()

# ── SPRINT 36 ────────────────────────────────────────────────────────────
sprint_header("C", 36, "Mobile PWA — Foundation", "Weeks 71–72",
"Build the Slipwise Mobile Progressive Web App: a responsive, installable web app with full navigation, task management, notifications, and a mobile-optimised layout.")

h2("10.13 Sprint 36 — Mobile PWA Foundation")

sprint_story(91, "Responsive Mobile Layout",
    "any Slipwise user accessing the app on a mobile device",
    "have a fully functional, mobile-optimised version of Slipwise that is fast and easy to use on a small screen",
    "I can stay productive from my phone without a degraded experience",
    13, "P0 — Must Have")
sprint_acceptance([
    "All pages in the Work OS, HRIS, and Notifications modules are fully responsive for screens 375px and wider.",
    "Mobile navigation uses a bottom tab bar with icons for: Home, Work OS, People, Notifications, More.",
    "The Task Detail Panel opens as a full-screen modal on mobile, not a side panel.",
    "All touch interactions are optimised: tap targets are minimum 44×44px, no hover-dependent features.",
    "The app passes Lighthouse Mobile audit with Performance ≥ 85, Accessibility ≥ 90.",
])

sprint_story(92, "PWA Installation and App Manifest",
    "any Slipwise user on mobile",
    "install Slipwise as an app on my home screen like a native app",
    "I have quick, distraction-free access to Slipwise without opening a browser",
    5, "P1 — Should Have")
sprint_acceptance([
    "web app manifest (manifest.json) configured with: name, short_name, icons (192×192, 512×512), theme_color, background_color, display: standalone.",
    "iOS Safari: 'Add to Home Screen' shows the Slipwise icon and opens in standalone mode.",
    "Android Chrome: 'Add to Home Screen' / 'Install App' prompt is shown when user visits 3+ times.",
    "Service Worker registered for PWA shell caching (Next.js routes, static assets).",
    "App shell loads instantly from cache even on slow connections.",
])
page_break()

# ── SPRINT 37 ────────────────────────────────────────────────────────────
sprint_header("C", 37, "Mobile — Offline Mode and Push Notifications", "Weeks 73–74",
"Add offline capability to the mobile PWA (read tasks, view profile, queue actions) and implement Web Push Notifications for task assignments, mentions, and ITSM alerts.")

h2("10.14 Sprint 37 — Mobile Offline and Push Notifications")

sprint_story(93, "Offline Mode",
    "a mobile user in an area with poor connectivity",
    "continue viewing my assigned tasks and queuing updates that sync when I am back online",
    "poor connectivity does not make me completely unable to work from my phone",
    8, "P2 — Could Have")
sprint_acceptance([
    "The Service Worker caches: the app shell, the current user's assigned tasks (last 100), the current user's profile, and the current leave balance.",
    "In offline mode, the app shows an 'Offline' banner. Cached data is displayed.",
    "Actions taken offline (task status change, comment draft) are stored in IndexedDB.",
    "On reconnection, queued actions are replayed via the API. Conflicts (another user changed the same task while offline) are shown with a merge prompt.",
])

sprint_story(94, "Web Push Notifications",
    "a mobile user",
    "receive push notifications on my phone for task assignments, mentions, leave approvals, and ITSM ticket updates",
    "I never miss urgent updates even when Slipwise is not open",
    8, "P1 — Should Have")
sprint_acceptance([
    "On PWA install or first visit, user is prompted to allow push notifications.",
    "Web Push uses the VAPID protocol. Push subscription is stored per user-device.",
    "Notification events: Task Assigned to Me, @mention in Comment, Leave Request Approved/Rejected, ITSM Ticket Assigned, Deal Won (for deal owner), CX Health Alert.",
    "Clicking a push notification deep-links to the relevant record in the app.",
    "Notification settings page allows disabling specific notification types.",
])
page_break()

# ── SPRINT 38 ────────────────────────────────────────────────────────────
sprint_header("C", 38, "Hardening, Security Audit, and Public Launch", "Weeks 75–76",
"The final sprint before public launch: full penetration testing, performance optimisation, database query analysis, accessibility audit, documentation completion, and go-live checklist execution.")

h2("10.15 Sprint 38 — Hardening, Audit, and Public Launch")
body("Sprint 38 is not a feature sprint. No new features are built. This sprint is entirely dedicated to ensuring the platform is production-ready, secure, performant, and polished. Every item in the launch checklist must be completed and signed off before the go-live date is confirmed.")

h3("Security Hardening Tasks")
security_tasks = [
    "Third-party penetration test (OWASP Top 10 scope) by external security firm. All Critical and High findings resolved.",
    "Complete audit of all API routes to verify checkPermission is called before data access.",
    "Review and tighten all Content Security Policy headers.",
    "Verify field-level encryption is applied to all salary and PII fields.",
    "Rotate all secrets in AWS Secrets Manager. Verify rotation policy is active.",
    "DDoS resilience test on the staging environment using a load testing tool.",
    "Verify all S3 buckets have Block Public Access enabled and bucket policies are restrictive.",
    "Confirm all Supabase RLS policies are active and tested for bypass scenarios.",
]
for t in security_tasks: bullet(t)

h3("Performance Optimisation Tasks")
perf_tasks = [
    "Analyse slow query log from RDS. Add missing indexes. Target: no query > 100ms on P95.",
    "Lighthouse audit all primary pages. Target: Performance ≥ 90 on desktop, ≥ 85 on mobile.",
    "Implement HTTP/2 push for critical CSS and JS bundles.",
    "Review and optimise all N+1 query patterns identified during QA.",
    "Configure CloudFront cache-control headers for all static assets (1 year TTL with content hash filenames).",
    "Verify Redis cache hit rate > 80% for permission checks under load.",
]
for t in perf_tasks: bullet(t)

h3("Accessibility Audit Tasks")
a11y_tasks = [
    "Full keyboard navigation audit: every action in the app must be keyboard-accessible.",
    "Screen reader compatibility test with NVDA (Windows) and VoiceOver (macOS/iOS).",
    "Colour contrast ratio must meet WCAG 2.1 AA standard for all text and UI elements.",
    "All images must have alt text. All icons must have aria-label attributes.",
    "Focus management in modals, panels, and dropdowns must be correct (focus trap in modals).",
]
for t in a11y_tasks: bullet(t)

h3("Go-Live Checklist")
golive = [
    "Production environment configured in AWS (all services, all secrets, all CDN rules).",
    "DNS cutover plan documented and rehearsed in staging.",
    "Database backup confirmed working with a test restore.",
    "Monitoring dashboards and alarms active in CloudWatch.",
    "On-call rotation schedule set up for the launch week.",
    "Customer support email and help docs published.",
    "Status page (status.slipwise.com) configured and embedded in the app.",
    "GDPR Privacy Policy and Terms of Service published and linked in the app footer.",
    "Post-launch communications (emails, social media announcements) scheduled.",
    "Rollback plan documented: if P1 issue detected within 24 hours of launch, process for reverting deployment.",
]
for t in golive: bullet(t)
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §11 — DATABASE SCHEMA MASTER REFERENCE
# ════════════════════════════════════════════════════════════════════════════
h1("11. Database Schema Master Reference")
body("This section provides the master reference for every database table in the Slipwise system. Tables are organised by module. Column types follow PostgreSQL conventions. All tables include id (cuid2), createdAt (timestamptz), and updatedAt (timestamptz) unless stated otherwise.")

db_modules = {
    "Platform / Identity": [
        ("Organisation", "id, name, slug(unique), logoUrl, timezone, currency, planId, status, createdAt"),
        ("User",          "id, email(unique), firstName, lastName, avatarUrl, supabaseAuthId, status, lastSeenAt"),
        ("OrganisationMember", "id, organisationId, userId, systemRole, roleTemplateId, employmentStatus, startDate, offboardedAt"),
        ("Invitation",    "id, organisationId, email, tokenHash, roleTemplateId, expiresAt, acceptedAt"),
    ],
    "RBAC": [
        ("RoleTemplate",  "id, organisationId, name, description, createdBy"),
        ("RoleTemplatePermission", "id, roleTemplateId, permissionKey, granted"),
        ("GranularPermissionGrant", "id, organisationId, userId, permissionKey, granted, grantedBy, grantedAt"),
    ],
    "Org Structure": [
        ("Department",    "id, organisationId, name, parentDepartmentId, headUserId, color, archivedAt"),
        ("Team",          "id, departmentId, name, description, teamLeadUserId, archivedAt"),
        ("TeamMembership","id, teamId, userId, joinedAt"),
    ],
    "Work OS": [
        ("Space",         "id, organisationId, name, description, color, icon, isPrivate, createdBy, archivedAt"),
        ("SpaceMember",   "id, spaceId, userId, role(ADMIN/MEMBER/VIEWER)"),
        ("CustomStatus",  "id, spaceId, name, color, type, sortOrder"),
        ("Folder",        "id, spaceId, name, color, sortOrder, archivedAt"),
        ("List",          "id, spaceId, folderId, name, description, color, defaultStatusId, sortOrder, archivedAt"),
        ("Sprint",        "id, listId, name, goal, startDate, endDate, status(PLANNED/ACTIVE/COMPLETED), velocity"),
        ("Task",          "id, listId, sprintId, parentTaskId, title, description(JSON), assigneeId, reporterId, statusId, priority, dueDate, startDate, timeEstimate, timeTracked, sortOrder, tags[], isRecurring, recurringConfig(JSON), archivedAt"),
        ("TaskDependency","id, predecessorTaskId, successorTaskId, type(FTS/STS/FTF)"),
        ("CustomFieldDef","id, spaceId, name, type, config(JSON), sortOrder"),
        ("TaskCustomFieldValue", "id, taskId, fieldDefinitionId, value(JSON)"),
        ("Comment",       "id, taskId, authorId, content(JSON), parentCommentId, editedAt"),
        ("Attachment",    "id, taskId, commentId, fileName, s3Key, mimeType, sizeBytes, uploadedBy"),
        ("TaskActivity",  "id, taskId, actorId, actionType, fieldName, oldValue, newValue, createdAt"),
        ("Form",          "id, listId, name, description, isPublic, fields(JSON), slug(unique)"),
        ("FormSubmission","id, formId, submitterEmail, data(JSON), taskId, createdAt"),
    ],
    "Time Tracking": [
        ("TimeEntry",     "id, taskId, userId, startTime, endTime, durationSeconds, note, type(TIMER/MANUAL)"),
    ],
    "HRIS": [
        ("EmployeeProfile", "id, userId, organisationId, personalEmail, workEmail, phone, pronouns, bio, timezone, jobTitle, employmentType, salaryBand(encrypted), noticeperiod, managerId, workAnniversary"),
        ("EquipmentRecord","id, employeeId, deviceType, serialNumber, assignedDate, returnedDate"),
        ("EmergencyContact","id, employeeId, name, relationship, phone"),
        ("LeaveType",     "id, organisationId, name, code, paidOrUnpaid, accrualType, annualAllowance, carryOverLimit, eligibleEmploymentTypes[]"),
        ("LeaveBalance",  "id, employeeId, leaveTypeId, year, totalAllowance, accrued, used, carryOver, available"),
        ("LeaveRequest",  "id, employeeId, leaveTypeId, startDate, endDate, workingDays, status, note, approvedBy, approvedAt, rejectionReason"),
        ("HolidayCalendar","id, organisationId, year, name, date, type(NATIONAL/COMPANY), country"),
        ("AttendanceRecord","id, employeeId, date, clockInTime, clockOutTime, totalMinutes, status(PRESENT/ABSENT/HALF_DAY/LEAVE), autoClosedAt"),
        ("ReviewCycle",   "id, organisationId, name, type, ratingScale, startDate, submissionDeadline, status"),
        ("ReviewAssignment","id, cycleId, revieweeId, reviewerId, reviewType(SELF/MANAGER/PEER), submittedAt"),
        ("ReviewResponse","id, assignmentId, competencyRatings(JSON), overallRating, summary"),
        ("CompensationRecord","id, employeeId, currency, grossSalary(encrypted), salaryBand, effectiveDate, paymentFrequency"),
        ("PayrollRun",    "id, organisationId, period, status, exportedAt, exportFormat"),
        ("PayrollRunItem","id, payrollRunId, employeeId, grossPay, deductions(JSON), netPay, workingDays, paidLeaveDays, unpaidLeaveDays"),
        ("OnboardingTemplate","id, organisationId, name, steps(JSON)"),
        ("OnboardingAssignment","id, templateId, employeeId, startedAt, completedAt, progress"),
    ],
    "CRM": [
        ("Contact",       "id, organisationId, firstName, lastName, email[], phone[], jobTitle, companyId, leadSource, contactStatus, leadScore, ownerId, tags[]"),
        ("Company",       "id, organisationId, name, domain, industry, size, annualRevenue, country, website, ownerId, tags[]"),
        ("Pipeline",      "id, organisationId, name, isDefault"),
        ("PipelineStage", "id, pipelineId, name, probability, rottenDays, color, sortOrder, isWon, isLost"),
        ("Deal",          "id, organisationId, title, pipelineId, stageId, contactId, companyId, ownerId, value, currency, probability, expectedCloseDate, actualCloseDate, lostReason, tags[]"),
        ("CRMActivity",   "id, organisationId, type, contactId, dealId, userId, subject, body, duration, outcome, scheduledAt, completedAt"),
        ("Sequence",      "id, organisationId, name, goal, status, steps(JSON)"),
        ("SequenceEnrolment","id, sequenceId, contactId, enrolledBy, status, currentStep, enrolledAt, completedAt"),
        ("LeadScoringModel","id, organisationId, name, rules(JSON)"),
        ("MeetingBooking","id, organisationId, hostUserId, contactId, title, startTime, endTime, zoomLink, calendarEventId"),
        ("SchedulingLink","id, userId, slug, title, duration, availableSlots(JSON)"),
    ],
    "ITSM": [
        ("ITSMService",   "id, organisationId, name, category, description, icon, assignedTeamId, defaultPriority, slaConfig(JSON), isDraft"),
        ("ITSMTicket",    "id, organisationId, serviceId, submitterId, assigneeId, title, description(JSON), priority, status, slaPolicy, firstResponseAt, resolvedAt, closedAt, csatScore, slaStatus"),
        ("TicketComment", "id, ticketId, authorId, content(JSON), isInternal, attachments[]"),
        ("RoutingRule",   "id, organisationId, serviceId, conditions(JSON), actions(JSON), priority, isActive"),
        ("EscalationPolicy","id, organisationId, serviceId, tiers(JSON), isActive"),
        ("KBArticle",     "id, organisationId, title, content(JSON), categoryId, authorId, status, viewCount, helpfulCount, notHelpfulCount"),
        ("KBCategory",    "id, organisationId, name, parentCategoryId, sortOrder"),
    ],
    "CX": [
        ("CXAccount",     "id, organisationId, companyId, csmOwnerId, accountStage, mrr, currency, contractRenewalDate, healthScore, healthTrend, lastHealthCalcAt"),
        ("CXHealthScoreBreakdown","id, accountId, usageScore, supportScore, npsScore, activityScore, renewalScore, calculatedAt"),
        ("CXPlaybook",    "id, organisationId, name, milestones(JSON)"),
        ("CXPlaybookAssignment","id, accountId, playbookId, assignedBy, startedAt, completedAt, progress"),
        ("NPS",           "id, organisationId, respondentEmail, score, comment, surveyDate, accountId"),
    ],
    "Goals / OKR": [
        ("Objective",     "id, organisationId, title, description, ownerId, ownerTeamId, parentObjectiveId, timeframe, year, status"),
        ("KeyResult",     "id, objectiveId, title, type, startValue, targetValue, currentValue, unit, dueDate"),
        ("KRCheckin",     "id, keyResultId, userId, previousValue, newValue, comment, createdAt"),
    ],
    "Docs": [
        ("Document",      "id, organisationId, spaceId, title, content(JSON/Yjs), parentDocumentId, createdBy, lastEditedBy, isPublic, publishedAt, archivedAt"),
        ("DocumentVersion","id, documentId, content(JSON), editedBy, wordCount, createdAt"),
    ],
    "Automation": [
        ("Automation",    "id, organisationId, name, triggerType, triggerConfig(JSON), nodes(JSON), isActive, lastRunAt"),
        ("AutomationRun", "id, automationId, triggeredAt, status(SUCCESS/ERROR/SKIPPED), nodeResults(JSON), errorMessage"),
        ("IncomingWebhook","id, organisationId, name, slug(unique), secretHash"),
        ("OutgoingWebhook","id, organisationId, url, events[], secretHash, isActive"),
    ],
    "Analytics": [
        ("Dashboard",     "id, organisationId, name, layout(JSON), isPublic, sharedWith[], createdBy"),
        ("DashboardWidget","id, dashboardId, type, dataSource, config(JSON), position(JSON)"),
        ("ScheduledReport","id, dashboardId, frequency, recipients[], lastSentAt, nextSendAt"),
    ],
    "Platform": [
        ("AuditLog",      "id, organisationId, actorId, actionType, entityType, entityId, oldValue(JSON), newValue(JSON), ipAddress, userAgent, createdAt"),
        ("Notification",  "id, recipientId, organisationId, type, title, body, entityType, entityId, readAt, createdAt"),
        ("CRMScoringModel","id, organisationId, name, rules(JSON), isActive"),
    ]
}

for module_name, tables in db_modules.items():
    h2(f"11.{list(db_modules.keys()).index(module_name)+1} {module_name}")
    header_table(
        ["Table Name", "Key Columns"],
        [[t, c] for t,c in tables],
        col_widths=[2.0, 5.1]
    )
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §12 — API ENDPOINT INDEX
# ════════════════════════════════════════════════════════════════════════════
h1("12. API Endpoint Index")
body("All Slipwise API routes are prefixed with /api/v1/. All routes require a valid Supabase JWT in the Authorization: Bearer {token} header. All state-mutating routes require the appropriate GranularPermissionGrant. All responses use JSON with a standard envelope: { success: boolean, data: T | null, error: string | null, meta: { page, pageSize, total } | null }.")

api_groups = [
    ("Organisation", [
        ("POST",   "/organisations",                    "Create organisation"),
        ("GET",    "/organisations/:id",                "Get organisation details"),
        ("PATCH",  "/organisations/:id",                "Update organisation settings"),
        ("GET",    "/organisations/:id/members",         "List all members"),
        ("POST",   "/members/invite",                   "Invite a new member"),
        ("PATCH",  "/members/:id",                      "Update member details"),
        ("DELETE", "/members/:id",                      "Deactivate member"),
        ("GET",    "/departments",                       "Get full department tree"),
        ("POST",   "/departments",                       "Create department"),
        ("PATCH",  "/departments/:id",                  "Update department"),
        ("DELETE", "/departments/:id",                  "Archive department"),
        ("GET",    "/teams",                             "List all teams"),
        ("POST",   "/teams",                             "Create team"),
        ("POST",   "/teams/:id/members",                "Add member to team"),
        ("DELETE", "/teams/:id/members/:userId",        "Remove member from team"),
    ]),
    ("RBAC", [
        ("GET",    "/role-templates",                   "List role templates"),
        ("POST",   "/role-templates",                   "Create role template"),
        ("PATCH",  "/role-templates/:id",               "Update role template"),
        ("GET",    "/permissions/users/:userId",        "Get user's permission grants"),
        ("POST",   "/permissions/grant",                "Create/update permission grant"),
        ("DELETE", "/permissions/grant/:id",            "Revoke permission grant"),
    ]),
    ("Work OS", [
        ("GET",    "/spaces",                           "List accessible spaces"),
        ("POST",   "/spaces",                           "Create space"),
        ("PATCH",  "/spaces/:id",                       "Update space"),
        ("DELETE", "/spaces/:id",                       "Archive space"),
        ("GET",    "/spaces/:id/statuses",              "Get custom statuses"),
        ("POST",   "/spaces/:id/statuses",              "Create custom status"),
        ("GET",    "/spaces/:id/lists",                 "List all lists in space"),
        ("POST",   "/lists",                            "Create list"),
        ("PATCH",  "/lists/:id",                        "Update list"),
        ("POST",   "/lists/:id/clone",                  "Clone list"),
        ("GET",    "/lists/:id/tasks",                  "Paginated task list with filters"),
        ("POST",   "/tasks",                            "Create task"),
        ("GET",    "/tasks/:id",                        "Get task detail"),
        ("PATCH",  "/tasks/:id",                        "Update task"),
        ("DELETE", "/tasks/:id",                        "Delete task"),
        ("POST",   "/tasks/:id/comments",               "Add comment"),
        ("PATCH",  "/tasks/:id/comments/:commentId",    "Edit comment"),
        ("DELETE", "/tasks/:id/comments/:commentId",    "Delete comment"),
        ("GET",    "/tasks/:id/activity",               "Get task activity log"),
        ("POST",   "/tasks/:id/dependencies",           "Add task dependency"),
        ("DELETE", "/tasks/:id/dependencies/:depId",    "Remove dependency"),
        ("GET",    "/sprints",                          "List sprints for a list"),
        ("POST",   "/sprints",                          "Create sprint"),
        ("POST",   "/sprints/:id/start",               "Start sprint"),
        ("POST",   "/sprints/:id/complete",            "Complete sprint with rollover options"),
        ("POST",   "/forms",                           "Create task intake form"),
        ("GET",    "/forms/:slug",                     "Get public form"),
        ("POST",   "/forms/:slug/submit",              "Submit form response"),
    ]),
    ("Time Tracking", [
        ("POST",   "/attendance/clock-in",              "Clock in"),
        ("POST",   "/attendance/clock-out",             "Clock out"),
        ("GET",    "/attendance",                       "Get attendance records (filtered)"),
        ("POST",   "/time-entries",                     "Manual time entry"),
        ("PATCH",  "/time-entries/:id",                 "Edit time entry"),
        ("DELETE", "/time-entries/:id",                 "Delete time entry"),
        ("GET",    "/reports/time-tracking",            "Time tracking report"),
    ]),
    ("HRIS", [
        ("GET",    "/employees/:id/profile",            "Get employee profile"),
        ("PATCH",  "/employees/:id/profile",            "Update employee profile"),
        ("GET",    "/leave-types",                      "List leave types"),
        ("POST",   "/leave-types",                      "Create leave type"),
        ("GET",    "/leave-balances/:employeeId",       "Get employee leave balances"),
        ("POST",   "/leave-requests",                   "Submit leave request"),
        ("PATCH",  "/leave-requests/:id/approve",       "Approve leave request"),
        ("PATCH",  "/leave-requests/:id/reject",        "Reject leave request"),
        ("PATCH",  "/leave-requests/:id/cancel",        "Cancel leave request"),
        ("GET",    "/holidays",                         "List holiday calendar"),
        ("POST",   "/holidays",                         "Add custom holiday"),
        ("GET",    "/review-cycles",                    "List review cycles"),
        ("POST",   "/review-cycles",                    "Create review cycle"),
        ("POST",   "/review-cycles/:id/activate",       "Activate review cycle"),
        ("GET",    "/review-assignments/:id",           "Get review assignment and form"),
        ("POST",   "/review-assignments/:id/submit",    "Submit review response"),
        ("GET",    "/payroll-runs",                     "List payroll runs"),
        ("POST",   "/payroll-runs",                     "Create payroll run"),
        ("POST",   "/payroll-runs/:id/finalise",        "Finalise payroll run"),
        ("GET",    "/payroll-runs/:id/export",          "Export payroll run"),
    ]),
    ("CRM", [
        ("GET",    "/contacts",                         "List contacts"),
        ("POST",   "/contacts",                         "Create contact"),
        ("GET",    "/contacts/:id",                     "Get contact detail"),
        ("PATCH",  "/contacts/:id",                     "Update contact"),
        ("DELETE", "/contacts/:id",                     "Delete contact"),
        ("POST",   "/contacts/import",                  "CSV import contacts"),
        ("GET",    "/companies",                        "List companies"),
        ("POST",   "/companies",                        "Create company"),
        ("GET",    "/companies/:id",                    "Get company detail"),
        ("GET",    "/pipelines",                        "List pipelines"),
        ("POST",   "/pipelines",                        "Create pipeline"),
        ("GET",    "/deals",                            "List deals (filterable)"),
        ("POST",   "/deals",                            "Create deal"),
        ("PATCH",  "/deals/:id",                        "Update deal"),
        ("PATCH",  "/deals/:id/stage",                  "Move deal stage"),
        ("POST",   "/crm/activities",                   "Log CRM activity"),
        ("GET",    "/crm/activities",                   "List activities (filtered)"),
        ("GET",    "/crm/sequences",                    "List sequences"),
        ("POST",   "/crm/sequences",                    "Create sequence"),
        ("POST",   "/crm/sequences/:id/enrol",          "Enrol contact in sequence"),
        ("GET",    "/crm/revenue",                      "Revenue dashboard data"),
    ]),
    ("ITSM", [
        ("GET",    "/itsm/services",                    "List service catalogue"),
        ("POST",   "/itsm/services",                    "Create service"),
        ("GET",    "/itsm/tickets",                     "List tickets (agent view)"),
        ("POST",   "/itsm/tickets",                     "Submit ticket"),
        ("GET",    "/itsm/tickets/:id",                 "Get ticket detail"),
        ("PATCH",  "/itsm/tickets/:id",                 "Update ticket"),
        ("POST",   "/itsm/tickets/:id/assign",          "Assign ticket to agent"),
        ("POST",   "/itsm/tickets/:id/resolve",         "Resolve ticket"),
        ("POST",   "/itsm/tickets/:id/close",           "Close ticket"),
        ("POST",   "/itsm/tickets/:id/comments",        "Add comment to ticket"),
        ("GET",    "/itsm/kb",                          "List KB articles"),
        ("POST",   "/itsm/kb",                          "Create KB article"),
        ("GET",    "/itsm/kb/:id",                      "Get KB article"),
        ("POST",   "/itsm/routing-rules",               "Create routing rule"),
        ("POST",   "/itsm/escalation-policies",         "Create escalation policy"),
    ]),
    ("CX", [
        ("GET",    "/cx/accounts",                      "List CX accounts"),
        ("POST",   "/cx/accounts",                      "Create CX account"),
        ("GET",    "/cx/accounts/:id",                  "Get CX account detail"),
        ("PATCH",  "/cx/accounts/:id",                  "Update CX account"),
        ("GET",    "/cx/playbooks",                     "List playbook templates"),
        ("POST",   "/cx/playbooks",                     "Create playbook template"),
        ("POST",   "/cx/accounts/:id/assign-playbook",  "Assign playbook to account"),
        ("GET",    "/cx/dashboard",                     "CX dashboard summary"),
    ]),
    ("Analytics", [
        ("GET",    "/analytics/dashboards",             "List dashboards"),
        ("POST",   "/analytics/dashboards",             "Create dashboard"),
        ("PATCH",  "/analytics/dashboards/:id",         "Update dashboard layout"),
        ("DELETE", "/analytics/dashboards/:id",         "Delete dashboard"),
        ("POST",   "/analytics/widgets",                "Create widget"),
        ("GET",    "/analytics/data/:source",           "Fetch data for widget"),
        ("POST",   "/analytics/reports/schedule",       "Schedule report delivery"),
    ]),
    ("AI", [
        ("POST",   "/ai/write",                         "AI writing / generation"),
        ("POST",   "/ai/summarise",                     "Summarise document or text"),
        ("POST",   "/ai/generate-tasks",                "Generate tasks from notes"),
        ("GET",    "/ai/insights",                      "Get AI insights for dashboard"),
    ]),
    ("Automation", [
        ("GET",    "/automation",                       "List automations"),
        ("POST",   "/automation",                       "Create automation"),
        ("PATCH",  "/automation/:id",                   "Update automation"),
        ("PATCH",  "/automation/:id/toggle",            "Enable / disable automation"),
        ("GET",    "/automation/:id/runs",              "Get execution log"),
        ("POST",   "/automation/webhooks",              "Create incoming webhook"),
        ("GET",    "/webhooks",                         "List outgoing webhooks"),
        ("POST",   "/webhooks",                         "Create outgoing webhook"),
    ]),
]

for group_name, endpoints in api_groups:
    h2(f"12.{list(e[0] for e in api_groups).index(group_name)+1} {group_name} Endpoints")
    header_table(
        ["Method", "Path", "Description"],
        [[m, p, d] for m,p,d in endpoints],
        col_widths=[0.7, 3.0, 3.4]
    )
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §13 — TESTING STRATEGY
# ════════════════════════════════════════════════════════════════════════════
h1("13. Testing Strategy")
body("Slipwise's testing strategy follows the testing pyramid: a broad base of unit tests, a middle layer of integration tests, and a narrow apex of end-to-end (E2E) tests for the most critical user journeys. The overall target is a code coverage of ≥ 75% for business logic functions, with 100% coverage for the permission check middleware, leave balance calculation, and SLA timer logic.")

h2("13.1 Unit Tests")
body("Unit tests cover individual functions and service-layer utilities in isolation. All external dependencies (database, Redis, email, AI APIs) are mocked. Unit test files live alongside their subject files with a .test.ts extension. The test runner is Vitest.")
bullet("Permission resolution function: tests for every combination of system role, role template, and individual grant/deny.")
bullet("Leave balance calculation: tests for FIXED_ANNUAL, MONTHLY_ACCRUAL, ANNIVERSARY accrual types including leap years.")
bullet("Working day calculation: tests including weekends, national holidays, custom holidays, and range spanning multiple months.")
bullet("Lead scoring model: tests for each rule type and combined score calculation.")
bullet("Health score formula: tests for each component weight and aggregate calculation.")
bullet("SLA time calculation: tests for business-hours-only SLA with various timezone configurations.")
bullet("Circular dependency detection for tasks: BFS traversal tests with chain, diamond, and cycle graph patterns.")
bullet("Automation rule engine: tests for each trigger type, condition evaluation, and action dispatch.")

h2("13.2 Integration Tests")
body("Integration tests run against a real PostgreSQL test database (seeded fresh before each test suite) and mock only external HTTP services (SES, S3, Gemini, Slack, Google). Integration tests verify that the full API handler stack works correctly end-to-end: request parsing → permission check → database operation → response formatting.")
bullet("Every POST endpoint: verify created record fields, verify permission check fires, verify 403 on missing permission.")
bullet("Multi-step workflows: leave request → approval → balance update → notification dispatch.")
bullet("Payroll run: create draft → auto-calculate items → finalise → export.")
bullet("Sprint completion: verify rollover logic for each option (backlog, next sprint, keep).")
bullet("Real-time: verify Supabase Realtime broadcast fires on task UPDATE.")

h2("13.3 End-to-End Tests")
body("E2E tests use Playwright to simulate real browser interactions. E2E tests run against the staging environment before every deployment to production.")

e2e_scenarios = [
    ("New Member Onboarding", "Admin invites → email received → member accepts → profile setup → appears in directory → correct permissions applied → can access allowed modules, blocked from disallowed."),
    ("Full Leave Lifecycle", "Employee submits leave → manager notified → manager approves → employee balance reduced → team calendar updated → employee cancels → balance restored."),
    ("Work OS Task Flow", "Create Space → Create List → Create Task → Assign to user → Change status (List View) → Move card (Board View) → Add comment with @mention → Mention notified."),
    ("CRM Deal to Task Automation", "Create Deal → Mark Won → Automation triggers → Work OS task list created from template → CS team notified in app."),
    ("ITSM Full Ticket Flow", "Browse service catalogue → Submit ticket → Ticket auto-routed → Agent assigned → Agent replies → SLA timer visible → Agent resolves → CSAT email sent → CSAT submitted → Score recorded."),
    ("Sprint Management", "Create list → Create sprint → Plan backlog → Move tasks to sprint → Start sprint → Complete tasks → End sprint → Incomplete tasks rolled to backlog → Velocity recorded."),
    ("Payroll Export", "Configure leave type → Employee takes approved leave → Month-end: create payroll run → Verify leave deduction in calculation → Finalise → Export CSV → Verify export file contents."),
    ("Permission Matrix", "Admin opens Permission Matrix for user → Toggles off 'crm:deals:view' → User navigates to CRM Deals → 403 error shown → Admin re-enables → User can access."),
]
header_table(
    ["Scenario", "Steps Covered"],
    e2e_scenarios,
    col_widths=[2.0, 5.1]
)
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §14 — MASTER DEVELOPMENT TIMELINE
# ════════════════════════════════════════════════════════════════════════════
h1("14. Master Development Timeline")
body("The master timeline below maps all 38 sprints to calendar dates, assuming a project start date of Week 1 = Month 1. Each sprint is two calendar weeks. Phase A (Sprints 1–9) is sequential — no sprint can begin until the previous passes QA. Phase B introduces parallelism: the CRM team (Sprints 17–20) runs in parallel with the ITSM team (Sprints 21–23) starting from Week 33. Phase C also has parallel tracks for AI (Sprints 32–33), Integrations (Sprints 34–35), and Mobile (Sprints 36–37).")

h2("14.1 Phase A — Sequential (Weeks 1–18)")
header_table(
    ["Sprint", "Name",                            "Start Week", "End Week", "Status Gate"],
    [
        ["01", "Org Tree and Entity Hierarchy",     "Week 1",  "Week 2",  "QA Pass + DB migration clean"],
        ["02", "RBAC Engine and Permission Matrix", "Week 3",  "Week 4",  "Security review sign-off"],
        ["03", "Member Onboarding Wizard",          "Week 5",  "Week 6",  "QA Pass + E2E: onboarding flow"],
        ["04", "Work OS Data Models",               "Week 7",  "Week 8",  "QA Pass + DB migration clean"],
        ["05", "Task CRUD and List View",           "Week 9",  "Week 10", "QA Pass + Perf: 200ms render"],
        ["06", "Board View, Dependencies, RT",      "Week 11", "Week 12", "QA Pass + RT: 2s update confirmed"],
        ["07", "Employee Profile and Directory",    "Week 13", "Week 14", "QA Pass"],
        ["08", "Leave Management Engine",           "Week 15", "Week 16", "QA Pass + Leave calc unit tests 100%"],
        ["09", "Attendance, Holidays, Calendar",    "Week 17", "Week 18", "QA Pass + Phase A staging deploy"],
    ],
    col_widths=[0.6, 2.8, 1.0, 1.0, 2.7]
)

h2("14.2 Phase B — Parallel Tracks (Weeks 19–46)")
body("Phase B runs three tracks simultaneously from Week 33 onward. Track 1 covers Work OS advanced features and HRIS depth. Track 2 covers CRM. Track 3 covers ITSM. All three tracks must merge and pass integrated QA before any Phase C sprint begins.")

header_table(
    ["Track",   "Sprint", "Name",                          "Weeks"],
    [
        ["B-Main",   "10", "Time Tracking",                 "19–20"],
        ["B-Main",   "11", "Sprint Management and Backlog", "21–22"],
        ["B-Main",   "12", "Gantt Chart and Mind Map",      "23–24"],
        ["B-Main",   "13", "Workload, Forms, Templates",    "25–26"],
        ["B-Main",   "14", "Performance Review Cycles",     "27–28"],
        ["B-Main",   "15", "Onboarding/Offboarding Auto",   "29–30"],
        ["B-Main",   "16", "Payroll Integration Layer",     "31–32"],
        ["B-CRM",    "17", "CRM Contacts and Companies",    "33–34"],
        ["B-CRM",    "18", "Pipelines, Deals, Revenue",     "35–36"],
        ["B-CRM",    "19", "Lead Scoring, Sequences",       "37–38"],
        ["B-CRM",    "20", "Meeting Scheduling, Activity",  "39–40"],
        ["B-ITSM",   "21", "ITSM Service Catalogue",        "41–42"],
        ["B-ITSM",   "22", "Ticket Routing, SLA, Escalation","43–44"],
        ["B-ITSM",   "23", "CSAT, Reporting, KB",           "45–46"],
    ],
    col_widths=[1.2, 0.6, 2.8, 2.5]
)

h2("14.3 Phase B → Phase C Integration QA (Weeks 47–48)")
body("Before Phase C begins, a dedicated two-week integration QA sprint runs all Phase B E2E tests, verifies cross-module data integrity, and ensures the combined staging environment passes the security scan and performance benchmarks.")

h2("14.4 Phase C — Intelligence and Scale (Weeks 47–76)")
header_table(
    ["Track",    "Sprint", "Name",                              "Weeks"],
    [
        ["C-Main",   "24", "CX Health Scoring, Playbooks",      "47–48"],
        ["C-Main",   "25", "Goals and OKR Engine",              "49–50"],
        ["C-Main",   "26", "Analytics Studio",                  "51–52"],
        ["C-Docs",   "27", "Docs and Wiki Foundation",          "53–54"],
        ["C-Docs",   "28", "Docs Embeds, Collab, Templates",    "55–56"],
        ["C-Auto",   "29", "Visual Automation Builder",         "57–58"],
        ["C-Auto",   "30", "Trigger Library, Integration Hooks","59–60"],
        ["C-Auto",   "31", "Zapier / Outgoing Webhooks",        "61–62"],
        ["C-AI",     "32", "AI Writing and Summarisation",      "63–64"],
        ["C-AI",     "33", "AI Predictive Analytics",           "65–66"],
        ["C-Int",    "34", "Google Workspace Integration",      "67–68"],
        ["C-Int",    "35", "Slack, Zoom, Calendar Sync",        "69–70"],
        ["C-Mobile", "36", "Mobile PWA Foundation",             "71–72"],
        ["C-Mobile", "37", "Mobile Offline, Push Notifs",       "73–74"],
        ["All",      "38", "Hardening, Audit, Launch",          "75–76"],
    ],
    col_widths=[1.2, 0.6, 2.8, 2.5]
)

h2("14.5 Resource Plan")
header_table(
    ["Phase", "Duration", "Team Composition",                            "Headcount"],
    [
        ["A", "18 weeks", "2 Backend + 2 Full-Stack + 1 QA + 0.5 DevOps", "5.5 FTE"],
        ["B (Main)", "28 weeks", "2 Backend + 2 Full-Stack + 1 QA",        "5 FTE"],
        ["B (CRM)",  "8 weeks",  "1 Backend + 2 Full-Stack",               "3 FTE"],
        ["B (ITSM)", "6 weeks",  "1 Backend + 1 Full-Stack + 0.5 QA",      "2.5 FTE"],
        ["C",        "30 weeks", "3 Full-Stack + 1 AI Eng + 1 Mobile + 1 QA + 0.5 DevOps", "7.5 FTE"],
        ["Total (concurrent peak)", "—", "All tracks running",             "~10 FTE"],
    ],
    col_widths=[1.5, 1.3, 3.5, 1.2]
)
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §15 — RISK REGISTER
# ════════════════════════════════════════════════════════════════════════════
h1("15. Risk Register")
body("The following risks have been identified during the planning phase. Each risk has an owner, a likelihood rating (H/M/L), an impact rating (H/M/L), and a mitigation strategy. The risk register should be reviewed at the start of every sprint by the engineering lead.")

header_table(
    ["Risk ID", "Risk Description",                               "Likelihood", "Impact", "Owner",  "Mitigation"],
    [
        ["R-01", "RBAC engine performance degrades under high user concurrency",    "M", "H", "Backend Lead", "Redis permission caching with 60s TTL; load test at Sprint 02 gate"],
        ["R-02", "Real-time Supabase Realtime channel limits reached for large orgs","M", "M", "Infra",        "Migrate to Redis Pub/Sub for orgs > 100 concurrent users; plan migration in Sprint 06"],
        ["R-03", "PostgreSQL query performance on complex filtered task queries",    "M", "H", "Backend Lead", "tsvector index on task title; quarterly EXPLAIN ANALYZE review"],
        ["R-04", "Gemini API quota limits during AI feature peak usage",            "M", "M", "AI Lead",      "Rate limiter on AI endpoints; implement request queuing; negotiate higher quota tiers"],
        ["R-05", "Third-party integration OAuth tokens expiring unexpectedly",      "L", "H", "Backend Lead", "Implement token refresh with pre-emptive refresh 1 hour before expiry"],
        ["R-06", "Payroll calculation errors due to timezone-related date bugs",    "M", "H", "HR Eng Lead",  "All date operations in UTC; convert to user timezone only at display layer; 100% unit test coverage for calculations"],
        ["R-07", "Phase B parallel teams creating conflicting database migrations", "H", "H", "Eng Lead",     "All migrations reviewed in weekly migration sync meeting; migration naming convention includes sprint number"],
        ["R-08", "Mobile PWA offline sync conflicts under poor connectivity",       "M", "M", "Mobile Lead",  "Use CRDT-based conflict resolution for offline queue; limit offline-editable entity types to tasks only in V1"],
        ["R-09", "GDPR data access requests requiring cross-module data export",    "L", "H", "Platform Eng", "Build a GDPR Data Export job that collects all user data across all tables; implement in Sprint 38 hardening"],
        ["R-10", "Scope creep from stakeholder requests during Phase B / C",       "H", "M", "Product Lead",  "All scope changes require a formal RFC document and sprint planning approval; no mid-sprint additions"],
    ],
    col_widths=[0.6, 2.5, 0.9, 0.7, 1.0, 2.4]
)
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  §16 — APPENDIX
# ════════════════════════════════════════════════════════════════════════════
h1("16. Appendix")
h2("16.1 Complete Permission Key Reference")
body("The following table lists every permission key defined in the system. All keys follow the format module:resource:action. Permission keys are defined in src/lib/permissions.ts as TypeScript string literal constants.")

permission_keys = [
    ("org:settings:view",           "View organisation settings"),
    ("org:settings:edit",           "Edit organisation name, logo, timezone, currency"),
    ("org:billing:view",            "View billing and subscription"),
    ("org:billing:manage",          "Change plan, update payment method"),
    ("org:members:invite",          "Invite new members"),
    ("org:members:deactivate",      "Deactivate members"),
    ("org:members:view",            "View member list and profiles"),
    ("org:permissions:manage",      "Grant / revoke permissions for others"),
    ("org:departments:manage",      "Create, edit, archive departments and teams"),
    ("org:org-chart:view",          "View org chart"),
    ("work-os:spaces:create",       "Create new Spaces"),
    ("work-os:spaces:manage",       "Edit or archive any Space"),
    ("work-os:tasks:create",        "Create tasks in accessible lists"),
    ("work-os:tasks:edit-own",      "Edit tasks created by self"),
    ("work-os:tasks:edit-any",      "Edit any task in accessible lists"),
    ("work-os:tasks:delete",        "Delete tasks"),
    ("work-os:tasks:assign",        "Assign tasks to other users"),
    ("work-os:sprints:manage",      "Create, start, and complete sprints"),
    ("work-os:time-tracking:log",   "Log time on tasks"),
    ("work-os:reports:view",        "View Work OS reports"),
    ("hris:profiles:view-own",      "View own employee profile"),
    ("hris:profiles:view-all",      "View all employee profiles"),
    ("hris:profiles:edit-own",      "Edit own profile public fields"),
    ("hris:profiles:edit-all",      "Edit any employee profile"),
    ("hris:salary:view",            "View salary and compensation data"),
    ("hris:leave:request",          "Submit leave requests"),
    ("hris:leave:approve",          "Approve or reject leave requests"),
    ("hris:leave:manage-types",     "Configure leave types and policies"),
    ("hris:attendance:view-own",    "View own attendance records"),
    ("hris:attendance:view-all",    "View all employees' attendance"),
    ("hris:attendance:edit",        "Edit attendance records (HR correction)"),
    ("hris:reviews:submit",         "Submit performance review responses"),
    ("hris:reviews:manage",         "Create and configure review cycles"),
    ("hris:payroll:view",           "View payroll run summaries"),
    ("hris:payroll:manage",         "Create, finalise, and export payroll runs"),
    ("hris:onboarding:manage",      "Configure onboarding/offboarding templates"),
    ("mailbox:view",                "View mailbox module (combined — per-mailbox grants layer on top)"),
    ("mailbox:{id}:read",           "Read specific mailbox (id = mailbox UUID)"),
    ("mailbox:{id}:send",           "Send from specific mailbox"),
    ("mailbox:{id}:manage",         "Manage mailbox settings"),
    ("docs:view",                   "View published documents"),
    ("docs:create",                 "Create new documents"),
    ("docs:edit-any",               "Edit any document"),
    ("docs:delete",                 "Delete documents"),
    ("docs:publish",                "Publish documents (make public)"),
    ("invoices:view",               "View invoices"),
    ("invoices:create",             "Create invoices"),
    ("invoices:send",               "Send invoices to clients"),
    ("invoices:delete",             "Delete draft invoices"),
    ("vouchers:view",               "View vouchers"),
    ("vouchers:create",             "Create vouchers"),
    ("vouchers:approve",            "Approve voucher claims"),
    ("crm:contacts:view",           "View CRM contacts"),
    ("crm:contacts:create",         "Create contacts"),
    ("crm:contacts:edit",           "Edit contacts"),
    ("crm:contacts:delete",         "Delete contacts"),
    ("crm:contacts:import",         "Import contacts via CSV"),
    ("crm:companies:view",          "View companies"),
    ("crm:companies:manage",        "Create, edit, delete companies"),
    ("crm:deals:view",              "View deals"),
    ("crm:deals:create",            "Create deals"),
    ("crm:deals:edit",              "Edit deals"),
    ("crm:deals:delete",            "Delete deals"),
    ("crm:pipelines:manage",        "Configure pipelines and stages"),
    ("crm:sequences:manage",        "Create and manage email sequences"),
    ("crm:revenue:view",            "View revenue dashboard"),
    ("itsm:tickets:view-own",       "View own submitted tickets"),
    ("itsm:tickets:view-all",       "View all tickets (agent view)"),
    ("itsm:tickets:create",         "Submit tickets"),
    ("itsm:tickets:assign",         "Assign tickets to agents"),
    ("itsm:tickets:resolve",        "Resolve and close tickets"),
    ("itsm:kb:view",                "View knowledge base articles"),
    ("itsm:kb:manage",              "Create and edit knowledge base articles"),
    ("itsm:settings:manage",        "Configure services, routing, SLA, escalation"),
    ("cx:accounts:view",            "View CX accounts"),
    ("cx:accounts:manage",          "Create and edit CX accounts"),
    ("cx:health:view",              "View health scores"),
    ("cx:playbooks:manage",         "Configure and assign playbooks"),
    ("analytics:dashboards:view",   "View shared dashboards"),
    ("analytics:dashboards:create", "Create personal and shared dashboards"),
    ("analytics:dashboards:manage", "Edit or delete any dashboard"),
    ("analytics:export",            "Export report data"),
    ("automation:view",             "View automations"),
    ("automation:create",           "Create automations"),
    ("automation:manage",           "Edit and delete any automation"),
    ("ai:write",                    "Use AI writing features"),
    ("ai:insights",                 "View AI insights feed"),
    ("admin:audit-log:view",        "View organisation audit log"),
    ("admin:webhooks:manage",       "Configure incoming and outgoing webhooks"),
    ("admin:integrations:manage",   "Connect and disconnect third-party integrations"),
]

header_table(
    ["Permission Key",                "Description"],
    [[k, d] for k,d in permission_keys],
    col_widths=[2.8, 4.3]
)

h2("16.2 Notification Types Reference")
notif_types = [
    ("TASK_ASSIGNED",         "Recipient: assignee. Trigger: task assigneeId changes to recipient."),
    ("TASK_COMMENT_MENTION",  "Recipient: mentioned user. Trigger: @mention in comment."),
    ("TASK_STATUS_CHANGED",   "Recipient: task watchers. Trigger: task status changes."),
    ("TASK_DUE_SOON",         "Recipient: assignee. Trigger: BullMQ job 24h before dueDate."),
    ("TASK_OVERDUE",          "Recipient: assignee + reporter. Trigger: BullMQ job at dueDate + 1 day."),
    ("LEAVE_REQUEST_SUBMITTED","Recipient: approver manager. Trigger: leave request created."),
    ("LEAVE_APPROVED",        "Recipient: employee. Trigger: leave request approved."),
    ("LEAVE_REJECTED",        "Recipient: employee. Trigger: leave request rejected."),
    ("REVIEW_FORM_ASSIGNED",  "Recipient: reviewer. Trigger: review cycle activated."),
    ("REVIEW_DEADLINE_REMINDER","Recipient: pending reviewers. Trigger: 48h before submission deadline."),
    ("DEAL_WON",              "Recipient: deal owner + CS team. Trigger: deal stage changes to Won stage."),
    ("DEAL_ROTTEN",           "Recipient: deal owner. Trigger: no activity for rottenDays."),
    ("ITSM_TICKET_ASSIGNED",  "Recipient: assigned agent. Trigger: ticket assignee set."),
    ("ITSM_TICKET_COMMENT",   "Recipient: submitter and watchers. Trigger: comment added."),
    ("ITSM_TICKET_RESOLVED",  "Recipient: submitter. Trigger: ticket status changes to RESOLVED."),
    ("ITSM_SLA_WARNING",      "Recipient: assigned agent. Trigger: 75% of SLA time elapsed."),
    ("ITSM_SLA_BREACH",       "Recipient: agent + team lead. Trigger: 100% SLA time elapsed."),
    ("CX_HEALTH_ALERT",       "Recipient: CSM owner. Trigger: health score drops ≥ 10 in one week."),
    ("OKR_CHECKIN_REMINDER",  "Recipient: KR owner. Trigger: Monday morning BullMQ cron."),
    ("AUTOMATION_ERROR",      "Recipient: automation creator. Trigger: automation run fails."),
    ("MEMBER_INVITATION",     "Recipient: invited email. Trigger: invite created."),
    ("CSAT_SURVEY",           "Recipient: ticket submitter. Trigger: 24h after ticket resolved."),
]
header_table(
    ["Notification Type", "Description"],
    [[t, d] for t,d in notif_types],
    col_widths=[2.2, 4.9]
)
page_break()

# ── Footer ───────────────────────────────────────────────────────────────
doc.add_paragraph()
divider()
p = doc.add_paragraph()
p.add_run("SLIPWISE ONE — Product Requirements Document v2.0").font.bold = True
doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y at %H:%M UTC')}")
doc.add_paragraph("Classification: INTERNAL — CONFIDENTIAL. Do not distribute externally.")
doc.add_paragraph("Office of the CTO, Slipwise Engineering Organisation.")

# ── Save ─────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"PRD generated successfully: {OUTPUT}")
